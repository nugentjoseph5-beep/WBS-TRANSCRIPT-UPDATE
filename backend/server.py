from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse, RedirectResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import asyncio
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import resend
from bson import ObjectId
import base64
import io
import httpx
from urllib.parse import urlencode

# Document generation imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET')
if not JWT_SECRET:
    raise ValueError("JWT_SECRET environment variable must be set")
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# Resend Configuration
RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', 'onboarding@resend.dev')
if RESEND_API_KEY:
    resend.api_key = RESEND_API_KEY

# Microsoft OAuth Configuration
MICROSOFT_CLIENT_ID = os.environ.get('MICROSOFT_CLIENT_ID', '')
MICROSOFT_CLIENT_SECRET = os.environ.get('MICROSOFT_CLIENT_SECRET', '')
MICROSOFT_TENANT_ID = os.environ.get('MICROSOFT_TENANT_ID', '')
MICROSOFT_REDIRECT_URI = os.environ.get('MICROSOFT_REDIRECT_URI', '')
GOOGLE_FORM_URL = os.environ.get('GOOGLE_FORM_URL', 'https://forms.gle/yCDU5p3EJ6X6yv7T9')
ALLOWED_EMAIL_DOMAIN = 'wolmers.org'

# Microsoft OAuth URLs
MICROSOFT_AUTH_URL = f"https://login.microsoftonline.com/{MICROSOFT_TENANT_ID}/oauth2/v2.0/authorize"
MICROSOFT_TOKEN_URL = f"https://login.microsoftonline.com/{MICROSOFT_TENANT_ID}/oauth2/v2.0/token"
MICROSOFT_GRAPH_URL = "https://graph.microsoft.com/v1.0/me"

# Create the main app
app = FastAPI(title="WBS Transcript and Recommendation Tracker API")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Upload directory
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# ==================== MODELS ====================

class UserBase(BaseModel):
    email: EmailStr
    full_name: str
    role: str = "student"  # student, staff, admin

class UserCreate(BaseModel):
    email: EmailStr
    full_name: str
    password: str
    role: str = "student"

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    full_name: str
    role: str
    created_at: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class TranscriptRequestCreate(BaseModel):
    first_name: str
    middle_name: Optional[str] = ""
    last_name: str
    date_of_birth: str  # mm/dd/yyyy format
    school_id: Optional[str] = ""
    enrollment_status: str  # enrolled, graduate, withdrawn
    academic_years: List[dict]  # List of {"from_year": "2015", "to_year": "2020"}
    wolmers_email: Optional[str] = ""
    personal_email: EmailStr
    phone_number: str
    last_form_class: Optional[str] = ""
    reason: str
    other_reason: Optional[str] = ""  # Required if reason is 'Other'
    needed_by_date: str
    collection_method: str  # pickup, emailed, delivery
    delivery_address: Optional[str] = ""  # Required if collection_method is 'delivery'
    institution_name: Optional[str] = ""
    institution_address: Optional[str] = ""
    institution_phone: Optional[str] = ""
    institution_email: Optional[str] = ""
    number_of_copies: Optional[int] = 1
    received_transcript_before: Optional[str] = ""  # YES or NO
    external_exams: Optional[List[dict]] = []  # List of {"exam": "CSEC", "year": "2020"} or {"exam": "Other", "name": "...", "year": "2021"}

class PasswordResetRequest(BaseModel):
    email: EmailStr

class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str

class AdministrativeClearance(BaseModel):
    no_fees_outstanding: Optional[bool] = False
    no_admin_obligations: Optional[bool] = False
    amount_paid: Optional[float] = None
    receipt_number: Optional[str] = ""
    payment_date: Optional[str] = ""
    updated_by: Optional[str] = ""
    updated_at: Optional[str] = ""

class TranscriptRequestUpdate(BaseModel):
    status: Optional[str] = None
    assigned_staff_id: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    note: Optional[str] = None  # Note for status changes
    administrative_clearance: Optional[AdministrativeClearance] = None

class TranscriptRequestResponse(BaseModel):
    id: str
    student_id: str
    student_name: str
    student_email: str
    first_name: str
    middle_name: str
    last_name: str
    date_of_birth: str = ""
    school_id: str = ""
    enrollment_status: str
    academic_years: List[dict] = []  # List of {"from_year": "2015", "to_year": "2020"}
    academic_year: str = ""  # Legacy field for backward compatibility
    wolmers_email: str = ""
    personal_email: str
    phone_number: str
    last_form_class: str = ""
    reason: str
    other_reason: str = ""
    needed_by_date: str
    collection_method: str
    delivery_address: str = ""
    institution_name: str = ""
    institution_address: str
    institution_phone: str
    institution_email: str
    number_of_copies: int = 1
    received_transcript_before: str = ""
    external_exams: List[dict] = []
    administrative_clearance: Optional[dict] = None
    status: str
    assigned_staff_id: Optional[str] = None
    assigned_staff_name: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    documents: List[dict] = []
    timeline: List[dict] = []
    created_at: str
    updated_at: str

class NotificationResponse(BaseModel):
    id: str
    user_id: str
    title: str
    message: str
    type: str
    read: bool
    request_id: Optional[str] = None
    created_at: str

class StaffCreateByAdmin(BaseModel):
    email: EmailStr
    full_name: str
    password: str
    role: str  # staff or admin

class AnalyticsResponse(BaseModel):
    total_requests: int
    pending_requests: int
    in_progress_requests: int
    processing_requests: int
    ready_requests: int
    completed_requests: int
    rejected_requests: int
    overdue_requests: int
    requests_by_month: List[dict]
    requests_by_enrollment: List[dict]
    recommendations_by_enrollment: List[dict]
    requests_by_collection_method: List[dict]
    staff_workload: List[dict]
    overdue_by_days: List[dict]
    # Recommendation letter stats
    total_recommendation_requests: int = 0
    pending_recommendation_requests: int = 0
    in_progress_recommendation_requests: int = 0
    completed_recommendation_requests: int = 0
    rejected_recommendation_requests: int = 0
    overdue_recommendation_requests: int = 0
    # Collection method breakdown for recommendations
    recommendations_by_collection_method: List[dict] = []
    # Overdue breakdown
    overdue_transcripts_by_days: List[dict] = []
    overdue_recommendations_by_days: List[dict] = []

# ==================== RECOMMENDATION LETTER MODELS ====================

class RecommendationRequestCreate(BaseModel):
    first_name: str
    middle_name: Optional[str] = ""
    last_name: str
    date_of_birth: str  # mm/dd/yyyy format
    email: EmailStr
    phone_number: str
    address: str
    years_attended: List[dict]  # List of {"from_year": "2015", "to_year": "2020"}
    enrollment_status: str  # Enrolled, Graduate, Withdrawn
    last_form_class: str  # e.g., "5W, 4R, 6AG2"
    co_curricular_activities: Optional[str] = ""  # Positions of responsibility and activities
    reason: str  # Reason for requesting recommendation
    other_reason: Optional[str] = ""  # Required if reason is 'Other'
    institution_name: str
    institution_address: str
    directed_to: Optional[str] = ""  # Whom should the letter be directed to
    program_name: str
    needed_by_date: str
    collection_method: str  # pickup, emailed, delivery
    delivery_address: Optional[str] = ""  # Required if collection_method is 'delivery'
    external_exams: Optional[List[dict]] = []  # List of {"exam": "CSEC", "year": "2020"} or {"exam": "Other", "name": "...", "year": "2021"}

class RecommendationRequestUpdate(BaseModel):
    status: Optional[str] = None
    assigned_staff_id: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    co_curricular_activities: Optional[str] = None
    note: Optional[str] = None  # Note for status changes
    administrative_clearance: Optional[dict] = None
    # Fields students can update for pending requests
    first_name: Optional[str] = None
    middle_name: Optional[str] = None
    last_name: Optional[str] = None
    email: Optional[str] = None
    phone_number: Optional[str] = None
    address: Optional[str] = None
    years_attended: Optional[List[dict]] = None
    enrollment_status: Optional[str] = None
    last_form_class: Optional[str] = None
    institution_name: Optional[str] = None
    institution_address: Optional[str] = None
    directed_to: Optional[str] = None
    program_name: Optional[str] = None
    needed_by_date: Optional[str] = None
    collection_method: Optional[str] = None
    delivery_address: Optional[str] = None
    external_exams: Optional[List[dict]] = None

class RecommendationRequestResponse(BaseModel):
    id: str
    student_id: str
    student_name: str
    student_email: str
    first_name: str
    middle_name: str
    last_name: str
    date_of_birth: str = ""
    email: str
    phone_number: str
    address: str
    years_attended: List[dict] = []  # List of {"from_year": "2015", "to_year": "2020"}
    years_attended_str: str = ""  # Legacy string format for backward compatibility
    enrollment_status: str = ""  # Added missing field
    last_form_class: str
    co_curricular_activities: str = ""
    reason: str = ""
    other_reason: str = ""
    institution_name: str
    institution_address: str
    directed_to: str
    program_name: str
    needed_by_date: str
    collection_method: str
    delivery_address: str = ""
    external_exams: List[dict] = []
    administrative_clearance: Optional[dict] = None
    status: str
    assigned_staff_id: Optional[str] = None
    assigned_staff_name: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    documents: List[dict] = []
    timeline: List[dict] = []
    created_at: str
    updated_at: str

class StudentRecommendationUpdate(BaseModel):
    first_name: Optional[str] = None
    middle_name: Optional[str] = None
    last_name: Optional[str] = None
    email: Optional[str] = None
    phone_number: Optional[str] = None
    address: Optional[str] = None
    years_attended: Optional[List[dict]] = None
    last_form_class: Optional[str] = None
    co_curricular_activities: Optional[str] = None
    institution_name: Optional[str] = None
    institution_address: Optional[str] = None
    directed_to: Optional[str] = None
    program_name: Optional[str] = None
    needed_by_date: Optional[str] = None
    collection_method: Optional[str] = None
    delivery_address: Optional[str] = None

# ==================== HELPER FUNCTIONS ====================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "role": role,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def decode_token(token: str) -> dict:
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    payload = decode_token(token)
    user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user

async def require_role(roles: List[str]):
    async def role_checker(user: dict = Depends(get_current_user)):
        if user["role"] not in roles:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return user
    return role_checker

async def send_email_notification(to_email: str, subject: str, html_content: str):
    if not RESEND_API_KEY:
        logger.warning("Resend API key not configured, skipping email")
        return None
    try:
        params = {
            "from": SENDER_EMAIL,
            "to": [to_email],
            "subject": subject,
            "html": html_content
        }
        result = await asyncio.to_thread(resend.Emails.send, params)
        logger.info(f"Email sent to {to_email}")
        return result
    except Exception as e:
        logger.error(f"Failed to send email: {str(e)}")
        return None

async def create_notification(user_id: str, title: str, message: str, notif_type: str, request_id: str = None):
    notification = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "title": title,
        "message": message,
        "type": notif_type,
        "read": False,
        "request_id": request_id,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(notification)
    return notification

async def check_and_notify_overdue_requests():
    """Check for overdue requests and notify admins"""
    now = datetime.now(timezone.utc)
    today_str = now.strftime("%Y-%m-%d")
    
    # Find overdue requests that haven't been notified today
    overdue_requests = await db.transcript_requests.find({
        "needed_by_date": {"$lt": today_str},
        "status": {"$nin": ["Completed", "Rejected"]},
        "$or": [
            {"overdue_notified_date": {"$exists": False}},
            {"overdue_notified_date": {"$ne": today_str}}
        ]
    }).to_list(None)
    
    if not overdue_requests:
        return
    
    # Get all admin users
    admins = await db.users.find({"role": "admin"}, {"_id": 0, "id": 1}).to_list(None)
    
    for req in overdue_requests:
        try:
            needed_date = datetime.strptime(req["needed_by_date"], "%Y-%m-%d")
            days_overdue = (now.replace(tzinfo=None) - needed_date).days
            
            student_name = f"{req.get('first_name', '')} {req.get('last_name', '')}"
            title = "⚠️ Overdue Transcript Request"
            message = f"Request from {student_name} is {days_overdue} day(s) overdue. Needed by: {req['needed_by_date']}"
            
            # Notify all admins
            for admin in admins:
                await create_notification(admin["id"], title, message, "overdue", req["id"])
            
            # Mark as notified today
            await db.transcript_requests.update_one(
                {"id": req["id"]},
                {"$set": {"overdue_notified_date": today_str}}
            )
        except Exception as e:
            print(f"Error notifying overdue request: {e}")

def normalize_recommendation_data(request_data: dict) -> dict:
    """Normalize recommendation request data for backward compatibility"""
    # Make a copy to avoid modifying original
    data = dict(request_data)
    
    # Handle years_attended field migration from string to list
    years_attended = data.get("years_attended", None)
    if years_attended is None:
        data["years_attended"] = []
        data["years_attended_str"] = ""
    elif isinstance(years_attended, str):
        # Convert old string format to new list format
        if years_attended and years_attended.strip() != "":
            # Handle formats like "2015-2020" or "2015-2020, 2021-2022"
            years_list = []
            for year_range in years_attended.split(", "):
                year_range = year_range.strip()
                if "-" in year_range:
                    parts = year_range.split("-", 1)
                    if len(parts) == 2:
                        years_list.append({"from_year": parts[0].strip(), "to_year": parts[1].strip()})
            data["years_attended"] = years_list
            data["years_attended_str"] = years_attended
        else:
            data["years_attended"] = []
            data["years_attended_str"] = ""
    elif isinstance(years_attended, list):
        # Already in new format, create string version for backward compatibility
        years_str = ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in years_attended if isinstance(y, dict)])
        data["years_attended_str"] = years_str
    else:
        data["years_attended"] = []
        data["years_attended_str"] = ""
    
    # Ensure all required fields have default values
    if "co_curricular_activities" not in data or data.get("co_curricular_activities") is None:
        data["co_curricular_activities"] = ""
    if "delivery_address" not in data or data.get("delivery_address") is None:
        data["delivery_address"] = ""
    if "years_attended_str" not in data:
        data["years_attended_str"] = ""
    if "reason" not in data or data.get("reason") is None:
        data["reason"] = ""
    if "other_reason" not in data or data.get("other_reason") is None:
        data["other_reason"] = ""
    if "enrollment_status" not in data or data.get("enrollment_status") is None:
        data["enrollment_status"] = ""
    if "date_of_birth" not in data or data.get("date_of_birth") is None:
        data["date_of_birth"] = ""
    if "external_exams" not in data or data.get("external_exams") is None:
        data["external_exams"] = []
    if "administrative_clearance" not in data or data.get("administrative_clearance") is None:
        data["administrative_clearance"] = None
    
    return data

def normalize_transcript_data(request_data: dict) -> dict:
    """Normalize transcript request data for backward compatibility"""
    # Make a copy to avoid modifying original
    data = dict(request_data)
    
    # Handle academic_years field migration from string to list
    academic_years = data.get("academic_years", None)
    if academic_years is None or (isinstance(academic_years, str)):
        # Convert old string format to new list format
        legacy_year = data.get("academic_year", "")
        if legacy_year and isinstance(legacy_year, str) and legacy_year.strip() != "":
            # Use the legacy academic_year field
            years_list = []
            for year_range in legacy_year.split(", "):
                year_range = year_range.strip()
                if "-" in year_range:
                    parts = year_range.split("-", 1)
                    if len(parts) == 2:
                        years_list.append({"from_year": parts[0].strip(), "to_year": parts[1].strip()})
            data["academic_years"] = years_list
        else:
            data["academic_years"] = []
    elif isinstance(academic_years, list):
        # Already in new format, create string version for backward compatibility
        years_str = ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in academic_years if isinstance(y, dict)])
        data["academic_year"] = years_str
    else:
        data["academic_years"] = []
    
    # Ensure academic_year legacy field exists
    if "academic_year" not in data or data.get("academic_year") is None:
        data["academic_year"] = ""
    
    # Ensure all required fields have default values
    if "delivery_address" not in data or data.get("delivery_address") is None:
        data["delivery_address"] = ""
    if "other_reason" not in data or data.get("other_reason") is None:
        data["other_reason"] = ""
    if "school_id" not in data or data.get("school_id") is None:
        data["school_id"] = ""
    if "wolmers_email" not in data or data.get("wolmers_email") is None:
        data["wolmers_email"] = ""
    if "date_of_birth" not in data or data.get("date_of_birth") is None:
        data["date_of_birth"] = ""
    if "last_form_class" not in data or data.get("last_form_class") is None:
        data["last_form_class"] = ""
    if "number_of_copies" not in data or data.get("number_of_copies") is None:
        data["number_of_copies"] = 1
    if "received_transcript_before" not in data or data.get("received_transcript_before") is None:
        data["received_transcript_before"] = ""
    if "external_exams" not in data or data.get("external_exams") is None:
        data["external_exams"] = []
    if "administrative_clearance" not in data or data.get("administrative_clearance") is None:
        data["administrative_clearance"] = None
    
    return data

def build_status_email_html(student_name: str, request_type: str, old_status: str, new_status: str,
                             request_id: str, rejection_reason: str = None, staff_notes: str = None,
                             accent_color: str = "#800000", app_url: str = "") -> str:
    """Build a rich HTML email for status change notifications."""
    status_color_map = {
        "Pending": "#eab308",
        "In Progress": "#3b82f6",
        "Processing": "#8b5cf6",
        "Ready": "#22c55e",
        "Completed": "#15803d",
        "Rejected": "#ef4444",
    }
    status_color = status_color_map.get(new_status, accent_color)
    detail_link = f"{app_url}/student" if app_url else ""
    rejection_block = ""
    if rejection_reason and new_status == "Rejected":
        rejection_block = f"""
        <div style="background-color:#fef2f2;border-left:4px solid #ef4444;padding:12px 16px;margin:12px 0;border-radius:4px;">
            <p style="margin:0;font-size:13px;color:#b91c1c;"><strong>Reason for Rejection:</strong> {rejection_reason}</p>
        </div>"""
    notes_block = ""
    if staff_notes:
        notes_block = f"""
        <div style="background-color:#f0f9ff;border-left:4px solid #0284c7;padding:12px 16px;margin:12px 0;border-radius:4px;">
            <p style="margin:0;font-size:13px;color:#0369a1;"><strong>Staff Notes:</strong> {staff_notes}</p>
        </div>"""
    return f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0"></head>
<body style="margin:0;padding:0;background-color:#f4f4f5;font-family:Arial,sans-serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="background-color:#f4f4f5;padding:30px 0;">
    <tr><td align="center">
      <table width="600" cellpadding="0" cellspacing="0" style="background:#ffffff;border-radius:8px;overflow:hidden;box-shadow:0 2px 8px rgba(0,0,0,0.08);max-width:600px;width:100%;">
        <!-- Header -->
        <tr><td style="background-color:{accent_color};padding:28px 32px;text-align:center;">
          <p style="margin:0;color:rgba(255,255,255,0.85);font-size:13px;letter-spacing:1px;text-transform:uppercase;">Wolmer&rsquo;s Boys&rsquo; School</p>
          <h1 style="margin:6px 0 0;color:#ffffff;font-size:22px;font-weight:700;">WBS Tracker</h1>
        </td></tr>
        <!-- Body -->
        <tr><td style="padding:32px;">
          <h2 style="margin:0 0 16px;font-size:18px;color:#1c1917;">{request_type} Status Update</h2>
          <p style="margin:0 0 20px;color:#57534e;font-size:14px;">Dear <strong>{student_name}</strong>,</p>
          <p style="margin:0 0 20px;color:#57534e;font-size:14px;">
            Your <strong>{request_type.lower()}</strong> request status has been updated. Here is a summary:
          </p>
          <!-- Status Change Box -->
          <table width="100%" cellpadding="0" cellspacing="0" style="background:#f9fafb;border:1px solid #e5e7eb;border-radius:8px;margin-bottom:20px;">
            <tr>
              <td style="padding:16px 20px;border-right:1px solid #e5e7eb;">
                <p style="margin:0;font-size:11px;color:#9ca3af;text-transform:uppercase;letter-spacing:0.5px;">Previous Status</p>
                <p style="margin:6px 0 0;font-size:15px;font-weight:600;color:#6b7280;">{old_status}</p>
              </td>
              <td style="padding:16px 20px;text-align:center;width:40px;color:#9ca3af;font-size:18px;">&rarr;</td>
              <td style="padding:16px 20px;">
                <p style="margin:0;font-size:11px;color:#9ca3af;text-transform:uppercase;letter-spacing:0.5px;">New Status</p>
                <p style="margin:6px 0 0;font-size:15px;font-weight:700;color:{status_color};">{new_status}</p>
              </td>
            </tr>
          </table>
          {rejection_block}
          {notes_block}
          {"<p style='margin:0 0 24px;'><a href='" + detail_link + "' style='display:inline-block;background-color:" + accent_color + ";color:#ffffff;padding:12px 24px;border-radius:6px;text-decoration:none;font-weight:600;font-size:14px;'>View My Requests</a></p>" if detail_link else ""}
        </td></tr>
        <!-- Footer -->
        <tr><td style="background-color:#f9fafb;padding:20px 32px;border-top:1px solid #e5e7eb;text-align:center;">
          <p style="margin:0;font-size:12px;color:#9ca3af;">
            This is an automated notification from the WBS Transcript &amp; Recommendation Tracker.<br>
            Wolmer&rsquo;s Boys&rsquo; School &bull; National Heroes Circle, Kingston 4
          </p>
        </td></tr>
      </table>
    </td></tr>
  </table>
</body>
</html>"""


async def notify_status_change(request_data: dict, old_status: str, new_status: str, rejection_reason: str = None, staff_notes: str = None):
    student = await db.users.find_one({"id": request_data["student_id"]}, {"_id": 0})
    if not student:
        return
    
    title = "Transcript Request Status Updated"
    message = f"Your transcript request has been updated from '{old_status}' to '{new_status}'."
    
    # Create in-app notification
    await create_notification(student["id"], title, message, "status_update", request_data["id"])
    
    # Determine email recipients: prefer wolmers_email, always include student email
    email_targets = set()
    wolmers_email = request_data.get("wolmers_email", "").strip()
    student_email = student.get("email", "").strip()
    personal_email = request_data.get("personal_email", "").strip()
    if wolmers_email:
        email_targets.add(wolmers_email)
    if student_email:
        email_targets.add(student_email)
    if personal_email:
        email_targets.add(personal_email)

    app_url = os.environ.get("APP_URL", MICROSOFT_REDIRECT_URI or "")
    html_content = build_status_email_html(
        student_name=student.get("full_name", "Student"),
        request_type="Transcript",
        old_status=old_status,
        new_status=new_status,
        request_id=request_data.get("id", ""),
        rejection_reason=rejection_reason,
        staff_notes=staff_notes,
        accent_color="#800000",
        app_url=app_url
    )
    subject = f"WBS Tracker – Transcript Request: {new_status}"
    for email in email_targets:
        await send_email_notification(email, subject, html_content)

# ==================== AUTH ROUTES ====================

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserCreate):
    # Check if email already exists
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Only students can self-register
    if user_data.role != "student":
        raise HTTPException(status_code=400, detail="Only students can self-register")
    
    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    user_doc = {
        "id": user_id,
        "email": user_data.email,
        "full_name": user_data.full_name,
        "password_hash": hash_password(user_data.password),
        "role": "student",
        "created_at": now,
        "updated_at": now
    }
    
    await db.users.insert_one(user_doc)
    
    token = create_token(user_id, user_data.email, "student")
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user_id,
            email=user_data.email,
            full_name=user_data.full_name,
            role="student",
            created_at=now
        )
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    token = create_token(user["id"], user["email"], user["role"])
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user["id"],
            email=user["email"],
            full_name=user["full_name"],
            role=user["role"],
            created_at=user["created_at"]
        )
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(user: dict = Depends(get_current_user)):
    return UserResponse(
        id=user["id"],
        email=user["email"],
        full_name=user["full_name"],
        role=user["role"],
        created_at=user["created_at"]
    )

# ==================== MICROSOFT OAUTH FOR STUDENTS ====================

class MicrosoftTokenRequest(BaseModel):
    code: str
    redirect_uri: str
    code_verifier: str  # PKCE code verifier

class MicrosoftAuthResponse(BaseModel):
    access_token: str
    user: UserResponse

@api_router.get("/auth/microsoft/config")
async def get_microsoft_config():
    """Get Microsoft OAuth configuration for frontend"""
    return {
        "client_id": MICROSOFT_CLIENT_ID,
        "tenant_id": MICROSOFT_TENANT_ID,
        "redirect_uri": MICROSOFT_REDIRECT_URI,
        "auth_url": MICROSOFT_AUTH_URL,
        "google_form_url": GOOGLE_FORM_URL,
        "allowed_domain": ALLOWED_EMAIL_DOMAIN
    }

@api_router.get("/auth/microsoft/login")
async def microsoft_login_redirect(code_challenge: str = None):
    """Redirect to Microsoft login page with PKCE support"""
    params = {
        "client_id": MICROSOFT_CLIENT_ID,
        "response_type": "code",
        "redirect_uri": MICROSOFT_REDIRECT_URI,
        "response_mode": "query",
        "scope": "openid profile email User.Read",
        "state": str(uuid.uuid4()),  # CSRF protection
    }
    
    # Add PKCE parameters if code_challenge is provided
    if code_challenge:
        params["code_challenge"] = code_challenge
        params["code_challenge_method"] = "S256"
    
    auth_url = f"{MICROSOFT_AUTH_URL}?{urlencode(params)}"
    return {"auth_url": auth_url}

@api_router.post("/auth/microsoft/callback")
async def microsoft_callback(request: MicrosoftTokenRequest):
    """Handle Microsoft OAuth callback and exchange code for tokens with PKCE"""
    try:
        # Exchange authorization code for access token with PKCE
        token_data = {
            "client_id": MICROSOFT_CLIENT_ID,
            "code": request.code,
            "redirect_uri": request.redirect_uri,
            "grant_type": "authorization_code",
            "scope": "openid profile email User.Read",
            "code_verifier": request.code_verifier  # PKCE code verifier
        }
        
        async with httpx.AsyncClient() as client:
            # Get access token from Microsoft
            token_response = await client.post(
                MICROSOFT_TOKEN_URL,
                data=token_data,
                headers={"Content-Type": "application/x-www-form-urlencoded"}
            )
            
            if token_response.status_code != 200:
                logger.error(f"Microsoft token error: {token_response.text}")
                raise HTTPException(status_code=400, detail="Failed to authenticate with Microsoft")
            
            tokens = token_response.json()
            access_token = tokens.get("access_token")
            
            if not access_token:
                raise HTTPException(status_code=400, detail="No access token received")
            
            # Get user info from Microsoft Graph API
            user_response = await client.get(
                MICROSOFT_GRAPH_URL,
                headers={"Authorization": f"Bearer {access_token}"}
            )
            
            if user_response.status_code != 200:
                logger.error(f"Microsoft Graph error: {user_response.text}")
                raise HTTPException(status_code=400, detail="Failed to get user info from Microsoft")
            
            user_info = user_response.json()
            
        # Extract email and validate domain
        email = user_info.get("mail") or user_info.get("userPrincipalName")
        display_name = user_info.get("displayName", "")
        
        if not email:
            raise HTTPException(status_code=400, detail="No email found in Microsoft account")
        
        # Validate email domain - CRITICAL SECURITY CHECK
        email_domain = email.split("@")[-1].lower()
        if email_domain != ALLOWED_EMAIL_DOMAIN.lower():
            logger.warning(f"Invalid email domain attempt: {email}")
            raise HTTPException(
                status_code=403, 
                detail=f"Only @{ALLOWED_EMAIL_DOMAIN} email addresses are permitted. Please use your Wolmer's Microsoft 365 account."
            )
        
        # Check if user exists or create new user
        existing_user = await db.users.find_one({"email": email.lower()}, {"_id": 0})
        now = datetime.now(timezone.utc).isoformat()
        
        if existing_user:
            # Update last login
            await db.users.update_one(
                {"email": email.lower()},
                {"$set": {"updated_at": now, "last_microsoft_login": now}}
            )
            user_id = existing_user["id"]
            full_name = existing_user["full_name"]
            role = existing_user["role"]
        else:
            # Create new student user
            user_id = str(uuid.uuid4())
            full_name = display_name or email.split("@")[0]
            role = "student"
            
            user_doc = {
                "id": user_id,
                "email": email.lower(),
                "full_name": full_name,
                "password_hash": "",  # Microsoft OAuth users don't have local password
                "role": role,
                "auth_provider": "microsoft",
                "microsoft_id": user_info.get("id"),
                "created_at": now,
                "updated_at": now,
                "last_microsoft_login": now
            }
            
            await db.users.insert_one(user_doc)
            logger.info(f"Created new Microsoft OAuth user: {email}")
        
        # Generate JWT token
        jwt_token = create_token(user_id, email.lower(), role)
        
        return MicrosoftAuthResponse(
            access_token=jwt_token,
            user=UserResponse(
                id=user_id,
                email=email.lower(),
                full_name=full_name,
                role=role,
                created_at=now
            )
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Microsoft OAuth error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Authentication failed: {str(e)}")

# ==================== MSAL.js CLIENT-SIDE AUTHENTICATION ====================

class MsalAuthRequest(BaseModel):
    """Request model for MSAL.js client-side authentication"""
    email: EmailStr
    name: str
    microsoft_id: str
    id_token: Optional[str] = None  # ID token from MSAL.js popup

@api_router.post("/auth/microsoft/authenticate", response_model=MicrosoftAuthResponse)
async def microsoft_authenticate_msal(request: MsalAuthRequest):
    """
    Authenticate user from MSAL.js client-side popup flow.
    This endpoint receives the user info directly from the frontend after 
    the MSAL.js library handles the OAuth popup authentication.
    """
    try:
        email = request.email.lower()
        
        # Validate email domain - CRITICAL SECURITY CHECK
        email_domain = email.split("@")[-1].lower()
        if email_domain != ALLOWED_EMAIL_DOMAIN.lower():
            logger.warning(f"Invalid email domain attempt: {email}")
            raise HTTPException(
                status_code=403, 
                detail=f"Only @{ALLOWED_EMAIL_DOMAIN} email addresses are permitted. Please use your Wolmer's Microsoft 365 account."
            )
        
        # Check if user exists or create new user
        existing_user = await db.users.find_one({"email": email}, {"_id": 0})
        now = datetime.now(timezone.utc).isoformat()
        
        if existing_user:
            # Update last login
            await db.users.update_one(
                {"email": email},
                {"$set": {"updated_at": now, "last_microsoft_login": now}}
            )
            user_id = existing_user["id"]
            full_name = existing_user["full_name"]
            role = existing_user["role"]
            created_at = existing_user.get("created_at", now)
            logger.info(f"Existing Microsoft OAuth user logged in: {email}")
        else:
            # Create new student user
            user_id = str(uuid.uuid4())
            full_name = request.name or email.split("@")[0]
            role = "student"
            created_at = now
            
            user_doc = {
                "id": user_id,
                "email": email,
                "full_name": full_name,
                "password_hash": "",  # Microsoft OAuth users don't have local password
                "role": role,
                "auth_provider": "microsoft",
                "microsoft_id": request.microsoft_id,
                "created_at": now,
                "updated_at": now,
                "last_microsoft_login": now
            }
            
            await db.users.insert_one(user_doc)
            logger.info(f"Created new Microsoft OAuth user: {email}")
        
        # Generate JWT token
        jwt_token = create_token(user_id, email, role)
        
        return MicrosoftAuthResponse(
            access_token=jwt_token,
            user=UserResponse(
                id=user_id,
                email=email,
                full_name=full_name,
                role=role,
                created_at=created_at
            )
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"MSAL authentication error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Authentication failed: {str(e)}")

# ==================== PASSWORD RESET ====================

@api_router.post("/auth/forgot-password")
async def forgot_password(request: PasswordResetRequest):
    user = await db.users.find_one({"email": request.email}, {"_id": 0})
    if not user:
        # Don't reveal if email exists or not for security
        return {"message": "If an account with this email exists, a password reset link has been sent."}
    
    # Generate reset token (valid for 1 hour)
    reset_token = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(hours=1)
    
    # Store reset token in database
    await db.password_resets.delete_many({"email": request.email})  # Remove old tokens
    await db.password_resets.insert_one({
        "token": reset_token,
        "email": request.email,
        "user_id": user["id"],
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Send email with reset link
    frontend_url = os.environ.get('FRONTEND_URL', 'http://localhost:3000')
    reset_link = f"{frontend_url}/reset-password?token={reset_token}"
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background-color: #800000; color: white; padding: 20px; text-align: center;">
            <h1 style="margin: 0;">Wolmer's Boys' School</h1>
            <p style="margin: 5px 0;">Transcript Tracker</p>
        </div>
        <div style="padding: 20px; background-color: #f5f5f5;">
            <h2 style="color: #800000;">Password Reset Request</h2>
            <p>Dear {user['full_name']},</p>
            <p>We received a request to reset your password. Click the button below to create a new password:</p>
            <div style="text-align: center; margin: 30px 0;">
                <a href="{reset_link}" style="background-color: #800000; color: white; padding: 12px 30px; text-decoration: none; border-radius: 5px; display: inline-block;">Reset Password</a>
            </div>
            <p style="color: #666; font-size: 14px;">This link will expire in 1 hour.</p>
            <p style="color: #666; font-size: 14px;">If you didn't request this, please ignore this email.</p>
            <p style="color: #999; font-size: 12px; margin-top: 30px;">
                This is an automated message from Wolmer's Boys' School Transcript Tracker.
            </p>
        </div>
    </div>
    """
    await send_email_notification(request.email, "Password Reset Request", html_content)
    
    # For development: log the token
    logger.info(f"Password reset token for {request.email}: {reset_token}")
    
    return {"message": "If an account with this email exists, a password reset link has been sent.", "token": reset_token}

@api_router.post("/auth/reset-password")
async def reset_password(request: PasswordResetConfirm):
    # Find the reset token
    reset_record = await db.password_resets.find_one({"token": request.token}, {"_id": 0})
    
    if not reset_record:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if token is expired
    expires_at = datetime.fromisoformat(reset_record["expires_at"])
    if datetime.now(timezone.utc) > expires_at:
        await db.password_resets.delete_one({"token": request.token})
        raise HTTPException(status_code=400, detail="Reset token has expired")
    
    # Update user's password
    new_password_hash = hash_password(request.new_password)
    await db.users.update_one(
        {"id": reset_record["user_id"]},
        {"$set": {"password_hash": new_password_hash, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Delete the used token
    await db.password_resets.delete_one({"token": request.token})
    
    return {"message": "Password has been reset successfully"}

@api_router.get("/auth/verify-reset-token/{token}")
async def verify_reset_token(token: str):
    reset_record = await db.password_resets.find_one({"token": token}, {"_id": 0})
    
    if not reset_record:
        raise HTTPException(status_code=400, detail="Invalid reset token")
    
    expires_at = datetime.fromisoformat(reset_record["expires_at"])
    if datetime.now(timezone.utc) > expires_at:
        await db.password_resets.delete_one({"token": token})
        raise HTTPException(status_code=400, detail="Reset token has expired")
    
    return {"valid": True, "email": reset_record["email"]}

# ==================== USER MANAGEMENT (ADMIN) ====================

@api_router.post("/admin/users", response_model=UserResponse)
async def create_user_by_admin(user_data: StaffCreateByAdmin, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create users")
    
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    user_doc = {
        "id": user_id,
        "email": user_data.email,
        "full_name": user_data.full_name,
        "password_hash": hash_password(user_data.password),
        "role": user_data.role,
        "created_at": now,
        "updated_at": now
    }
    
    await db.users.insert_one(user_doc)
    
    return UserResponse(
        id=user_id,
        email=user_data.email,
        full_name=user_data.full_name,
        role=user_data.role,
        created_at=now
    )

@api_router.get("/admin/users", response_model=List[UserResponse])
async def get_all_users(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view all users")
    
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).to_list(1000)
    return [UserResponse(**u) for u in users]

@api_router.get("/admin/staff", response_model=List[UserResponse])
async def get_staff_members(current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    staff = await db.users.find({"role": "staff"}, {"_id": 0, "password_hash": 0}).to_list(1000)
    return [UserResponse(**s) for s in staff]

@api_router.delete("/admin/users/{user_id}")
async def delete_user(user_id: str, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete users")
    
    if user_id == current_user["id"]:
        raise HTTPException(status_code=400, detail="Cannot delete yourself")
    
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "User deleted successfully"}

class AdminResetPassword(BaseModel):
    new_password: str

@api_router.post("/admin/users/{user_id}/reset-password")
async def admin_reset_user_password(user_id: str, data: AdminResetPassword, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can reset passwords")
    
    if user_id == current_user["id"]:
        raise HTTPException(status_code=400, detail="Cannot reset your own password this way. Use the forgot password feature.")
    
    # Find the user
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Update password
    new_password_hash = hash_password(data.new_password)
    await db.users.update_one(
        {"id": user_id},
        {"$set": {"password_hash": new_password_hash, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Send notification email to user
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background-color: #800000; color: white; padding: 20px; text-align: center;">
            <h1 style="margin: 0;">Wolmer's Boys' School</h1>
            <p style="margin: 5px 0;">Transcript Tracker</p>
        </div>
        <div style="padding: 20px; background-color: #f5f5f5;">
            <h2 style="color: #800000;">Password Reset by Administrator</h2>
            <p>Dear {user['full_name']},</p>
            <p>Your password has been reset by an administrator.</p>
            <p>If you did not request this change, please contact the administrator immediately.</p>
            <p style="color: #999; font-size: 12px; margin-top: 30px;">
                This is an automated message from Wolmer's Boys' School Transcript Tracker.
            </p>
        </div>
    </div>
    """
    await send_email_notification(user["email"], "Your Password Has Been Reset", html_content)
    
    return {"message": f"Password reset successfully for {user['full_name']}"}

# ==================== TRANSCRIPT REQUESTS ====================

@api_router.post("/requests", response_model=TranscriptRequestResponse)
async def create_transcript_request(request_data: TranscriptRequestCreate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can create transcript requests")
    
    request_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request submitted",
        "updated_by": current_user["full_name"]
    }
    
    # Format academic years for display (backward compatibility)
    academic_years_str = ", ".join([f"{y['from_year']}-{y['to_year']}" for y in request_data.academic_years]) if request_data.academic_years else ""
    
    doc = {
        "id": request_id,
        "student_id": current_user["id"],
        "student_name": current_user["full_name"],
        "student_email": current_user["email"],
        "first_name": request_data.first_name,
        "middle_name": request_data.middle_name or "",
        "last_name": request_data.last_name,
        "date_of_birth": request_data.date_of_birth,
        "school_id": request_data.school_id or "",
        "enrollment_status": request_data.enrollment_status,
        "academic_years": request_data.academic_years,
        "academic_year": academic_years_str,  # Legacy field
        "wolmers_email": request_data.wolmers_email or "",
        "personal_email": request_data.personal_email,
        "phone_number": request_data.phone_number,
        "last_form_class": request_data.last_form_class or "",
        "reason": request_data.reason,
        "other_reason": request_data.other_reason or "",
        "needed_by_date": request_data.needed_by_date,
        "collection_method": request_data.collection_method,
        "delivery_address": request_data.delivery_address or "",
        "institution_name": request_data.institution_name or "",
        "institution_address": request_data.institution_address or "",
        "institution_phone": request_data.institution_phone or "",
        "institution_email": request_data.institution_email or "",
        "number_of_copies": request_data.number_of_copies or 1,
        "received_transcript_before": request_data.received_transcript_before or "",
        "external_exams": request_data.external_exams or [],
        "status": "Pending",
        "assigned_staff_id": None,
        "assigned_staff_name": None,
        "rejection_reason": None,
        "staff_notes": None,
        "documents": [],
        "timeline": [timeline_entry],
        "created_at": now,
        "updated_at": now
    }
    
    await db.transcript_requests.insert_one(doc)
    
    # Notify admins
    admins = await db.users.find({"role": "admin"}, {"_id": 0}).to_list(100)
    for admin in admins:
        await create_notification(
            admin["id"],
            "New Transcript Request",
            f"New request from {current_user['full_name']}",
            "new_request",
            request_id
        )
    
    return TranscriptRequestResponse(**doc)

@api_router.get("/requests", response_model=List[TranscriptRequestResponse])
async def get_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] == "student":
        # Students can only see their own requests
        requests = await db.transcript_requests.find(
            {"student_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    elif current_user["role"] == "staff":
        # Staff can see assigned requests
        requests = await db.transcript_requests.find(
            {"assigned_staff_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    else:
        # Admin can see all requests
        requests = await db.transcript_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Normalize data for backward compatibility
    normalized_requests = [normalize_transcript_data(r) for r in requests]
    return [TranscriptRequestResponse(**r) for r in normalized_requests]

@api_router.get("/requests/all", response_model=List[TranscriptRequestResponse])
async def get_all_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    requests = await db.transcript_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    # Normalize data for backward compatibility
    normalized_requests = [normalize_transcript_data(r) for r in requests]
    return [TranscriptRequestResponse(**r) for r in normalized_requests]

@api_router.get("/requests/{request_id}", response_model=TranscriptRequestResponse)
async def get_request(request_id: str, current_user: dict = Depends(get_current_user)):
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check permissions
    if current_user["role"] == "student" and request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only view your own requests")
    
    # Normalize data for backward compatibility
    normalized_request = normalize_transcript_data(request_doc)
    return TranscriptRequestResponse(**normalized_request)

class StudentRequestUpdate(BaseModel):
    first_name: Optional[str] = None
    middle_name: Optional[str] = None
    last_name: Optional[str] = None
    school_id: Optional[str] = None
    enrollment_status: Optional[str] = None
    academic_year: Optional[str] = None
    wolmers_email: Optional[str] = None
    personal_email: Optional[str] = None
    phone_number: Optional[str] = None
    reason: Optional[str] = None
    needed_by_date: Optional[str] = None
    collection_method: Optional[str] = None
    institution_name: Optional[str] = None
    institution_address: Optional[str] = None
    institution_phone: Optional[str] = None
    institution_email: Optional[str] = None

@api_router.put("/requests/{request_id}/edit", response_model=TranscriptRequestResponse)
async def student_edit_request(request_id: str, update_data: StudentRequestUpdate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can edit their own requests")
    
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check ownership
    if request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only edit your own requests")
    
    # Check status - only pending requests can be edited
    if request_doc["status"] != "Pending":
        raise HTTPException(
            status_code=400, 
            detail=f"This request cannot be edited because its status is '{request_doc['status']}'. Only pending requests can be modified."
        )
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    # Update only provided fields
    update_fields = update_data.dict(exclude_unset=True)
    for field, value in update_fields.items():
        if value is not None:
            updates[field] = value
    
    # Add timeline entry for edit
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request details updated by student",
        "updated_by": current_user["full_name"]
    }
    
    await db.transcript_requests.update_one(
        {"id": request_id},
        {
            "$set": updates,
            "$push": {"timeline": timeline_entry}
        }
    )
    
    updated_request = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    return TranscriptRequestResponse(**updated_request)

@api_router.patch("/requests/{request_id}", response_model=TranscriptRequestResponse)
async def update_request(request_id: str, update_data: TranscriptRequestUpdate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] == "student":
        raise HTTPException(status_code=403, detail="Students cannot update request status")
    
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    old_status = request_doc["status"]
    
    if update_data.status:
        updates["status"] = update_data.status
        # Use the provided note or create a default one
        note_text = update_data.note if update_data.note else f"Status changed to {update_data.status}"
        timeline_entry = {
            "status": update_data.status,
            "timestamp": now,
            "note": note_text,
            "updated_by": current_user["full_name"]
        }
        await db.transcript_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.assigned_staff_id:
        staff = await db.users.find_one({"id": update_data.assigned_staff_id}, {"_id": 0})
        if staff:
            updates["assigned_staff_id"] = update_data.assigned_staff_id
            updates["assigned_staff_name"] = staff["full_name"]
            
            # Auto-update status to "In Progress" when staff is assigned
            if request_doc.get("status") == "Pending":
                updates["status"] = "In Progress"
                timeline_entry = {
                    "status": "In Progress",
                    "timestamp": now,
                    "note": "Request assigned to staff - Status automatically updated to In Progress",
                    "updated_by": current_user["full_name"]
                }
                await db.transcript_requests.update_one(
                    {"id": request_id},
                    {"$push": {"timeline": timeline_entry}}
                )
            
            # Notify staff
            await create_notification(
                staff["id"],
                "New Assignment",
                f"You have been assigned a transcript request",
                "assignment",
                request_id
            )
    
    if update_data.rejection_reason:
        updates["rejection_reason"] = update_data.rejection_reason
        updates["status"] = "Rejected"
        timeline_entry = {
            "status": "Rejected",
            "timestamp": now,
            "note": f"Request rejected: {update_data.rejection_reason}",
            "updated_by": current_user["full_name"]
        }
        await db.transcript_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.staff_notes:
        updates["staff_notes"] = update_data.staff_notes
    
    if update_data.administrative_clearance is not None:
        updates["administrative_clearance"] = update_data.administrative_clearance.model_dump()
    
    await db.transcript_requests.update_one({"id": request_id}, {"$set": updates})
    
    # Notify student of status change
    effective_new_status = update_data.status or ("Rejected" if update_data.rejection_reason else None)
    if effective_new_status and effective_new_status != old_status:
        updated_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
        await notify_status_change(
            updated_doc, old_status, effective_new_status,
            rejection_reason=update_data.rejection_reason,
            staff_notes=update_data.staff_notes
        )
    
    updated_request = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    return TranscriptRequestResponse(**updated_request)

# ==================== FILE UPLOAD ====================

@api_router.post("/requests/{request_id}/documents")
async def upload_document(request_id: str, file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Only staff and admin can upload documents")
    
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Validate file type
    allowed_types = [
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "image/jpeg",
        "image/png",
        "image/gif"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not allowed")
    
    # Save file
    file_id = str(uuid.uuid4())
    file_ext = Path(file.filename).suffix
    file_path = UPLOAD_DIR / f"{file_id}{file_ext}"
    
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    now = datetime.now(timezone.utc).isoformat()
    doc_entry = {
        "id": file_id,
        "filename": file.filename,
        "content_type": file.content_type,
        "path": str(file_path),
        "uploaded_by": current_user["full_name"],
        "uploaded_at": now
    }
    
    await db.transcript_requests.update_one(
        {"id": request_id},
        {
            "$push": {"documents": doc_entry},
            "$set": {"updated_at": now}
        }
    )
    
    # Add timeline entry
    timeline_entry = {
        "status": request_doc["status"],
        "timestamp": now,
        "note": f"Document uploaded: {file.filename}",
        "updated_by": current_user["full_name"]
    }
    await db.transcript_requests.update_one(
        {"id": request_id},
        {"$push": {"timeline": timeline_entry}}
    )
    
    # Notify student
    await create_notification(
        request_doc["student_id"],
        "Document Uploaded",
        f"A document has been uploaded to your transcript request",
        "document",
        request_id
    )
    
    return {"message": "Document uploaded successfully", "document": doc_entry}

@api_router.get("/documents/{document_id}")
async def get_document(document_id: str, current_user: dict = Depends(get_current_user)):
    # Find the request containing this document
    request_doc = await db.transcript_requests.find_one(
        {"documents.id": document_id},
        {"_id": 0}
    )
    
    if not request_doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check permissions
    if current_user["role"] == "student" and request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Find the document
    doc = next((d for d in request_doc["documents"] if d["id"] == document_id), None)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = Path(doc["path"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on server")
    
    with open(file_path, "rb") as f:
        content = base64.b64encode(f.read()).decode('utf-8')
    
    return {
        "filename": doc["filename"],
        "content_type": doc["content_type"],
        "content": content
    }

# ==================== RECOMMENDATION LETTER REQUESTS ====================

@api_router.post("/recommendations", response_model=RecommendationRequestResponse)
async def create_recommendation_request(request_data: RecommendationRequestCreate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can create recommendation letter requests")
    
    request_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request submitted",
        "updated_by": current_user["full_name"]
    }
    
    # Format years attended for display (backward compatibility)
    years_attended_str = ", ".join([f"{y['from_year']}-{y['to_year']}" for y in request_data.years_attended]) if request_data.years_attended else ""
    
    doc = {
        "id": request_id,
        "student_id": current_user["id"],
        "student_name": current_user["full_name"],
        "student_email": current_user["email"],
        "first_name": request_data.first_name,
        "middle_name": request_data.middle_name or "",
        "last_name": request_data.last_name,
        "date_of_birth": request_data.date_of_birth,
        "email": request_data.email,
        "phone_number": request_data.phone_number,
        "address": request_data.address,
        "years_attended": request_data.years_attended,
        "years_attended_str": years_attended_str,  # Legacy field
        "enrollment_status": request_data.enrollment_status,  # Added missing field
        "last_form_class": request_data.last_form_class,
        "co_curricular_activities": request_data.co_curricular_activities or "",
        "reason": request_data.reason,
        "other_reason": request_data.other_reason or "",
        "institution_name": request_data.institution_name,
        "institution_address": request_data.institution_address,
        "directed_to": request_data.directed_to or "",
        "program_name": request_data.program_name,
        "needed_by_date": request_data.needed_by_date,
        "collection_method": request_data.collection_method,
        "delivery_address": request_data.delivery_address or "",
        "external_exams": request_data.external_exams or [],
        "status": "Pending",
        "assigned_staff_id": None,
        "assigned_staff_name": None,
        "rejection_reason": None,
        "staff_notes": None,
        "documents": [],
        "timeline": [timeline_entry],
        "created_at": now,
        "updated_at": now
    }
    
    await db.recommendation_requests.insert_one(doc)
    
    # Notify admins
    admins = await db.users.find({"role": "admin"}, {"_id": 0}).to_list(100)
    for admin in admins:
        await create_notification(
            admin["id"],
            "New Recommendation Letter Request",
            f"New recommendation letter request from {current_user['full_name']}",
            "new_recommendation",
            request_id
        )
    
    return RecommendationRequestResponse(**doc)

@api_router.get("/recommendations", response_model=List[RecommendationRequestResponse])
async def get_recommendation_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] == "student":
        # Students can only see their own requests
        requests = await db.recommendation_requests.find(
            {"student_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    elif current_user["role"] == "staff":
        # Staff can see assigned requests
        requests = await db.recommendation_requests.find(
            {"assigned_staff_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    else:
        # Admin can see all requests
        requests = await db.recommendation_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Normalize data for backward compatibility
    normalized_requests = [normalize_recommendation_data(r) for r in requests]
    return [RecommendationRequestResponse(**r) for r in normalized_requests]

@api_router.get("/recommendations/all", response_model=List[RecommendationRequestResponse])
async def get_all_recommendation_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    requests = await db.recommendation_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    # Normalize data for backward compatibility
    normalized_requests = [normalize_recommendation_data(r) for r in requests]
    return [RecommendationRequestResponse(**r) for r in normalized_requests]

@api_router.get("/recommendations/{request_id}", response_model=RecommendationRequestResponse)
async def get_recommendation_request(request_id: str, current_user: dict = Depends(get_current_user)):
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check permissions
    if current_user["role"] == "student" and request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only view your own requests")
    
    # Normalize data for backward compatibility
    normalized_request = normalize_recommendation_data(request_doc)
    return RecommendationRequestResponse(**normalized_request)

@api_router.put("/recommendations/{request_id}/edit", response_model=RecommendationRequestResponse)
async def student_edit_recommendation(request_id: str, update_data: StudentRecommendationUpdate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can edit their own requests")
    
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check ownership
    if request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only edit your own requests")
    
    # Check status - only pending requests can be edited
    if request_doc["status"] != "Pending":
        raise HTTPException(
            status_code=400, 
            detail=f"This request cannot be edited because its status is '{request_doc['status']}'. Only pending requests can be modified."
        )
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    # Update only provided fields
    update_fields = update_data.dict(exclude_unset=True)
    for field, value in update_fields.items():
        if value is not None:
            updates[field] = value
    
    # Add timeline entry for edit
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request details updated by student",
        "updated_by": current_user["full_name"]
    }
    
    await db.recommendation_requests.update_one(
        {"id": request_id},
        {
            "$set": updates,
            "$push": {"timeline": timeline_entry}
        }
    )
    
    updated_request = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    normalized_request = normalize_recommendation_data(updated_request)
    return RecommendationRequestResponse(**normalized_request)

@api_router.patch("/recommendations/{request_id}", response_model=RecommendationRequestResponse)
async def update_recommendation_request(request_id: str, update_data: RecommendationRequestUpdate, current_user: dict = Depends(get_current_user)):
    # Allow students to update their own pending recommendations
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check permissions
    if current_user["role"] == "student":
        # Students can only update their own requests and only if status is Pending or In Progress
        if request_doc["student_id"] != current_user["id"]:
            raise HTTPException(status_code=403, detail="You can only update your own requests")
        if request_doc["status"] not in ["Pending", "In Progress"]:
            raise HTTPException(status_code=403, detail=f"You cannot edit requests at the '{request_doc['status']}' stage. Please contact administration for assistance.")
        # Students cannot change status, assign staff, or reject
        if update_data.status or update_data.assigned_staff_id or update_data.rejection_reason:
            raise HTTPException(status_code=403, detail="Students cannot update request status or assignments")
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    old_status = request_doc["status"]
    
    if update_data.status:
        updates["status"] = update_data.status
        # Use the provided note or create a default one
        note_text = update_data.note if update_data.note else f"Status changed to {update_data.status}"
        timeline_entry = {
            "status": update_data.status,
            "timestamp": now,
            "note": note_text,
            "updated_by": current_user["full_name"]
        }
        await db.recommendation_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.assigned_staff_id:
        staff = await db.users.find_one({"id": update_data.assigned_staff_id}, {"_id": 0})
        if staff:
            updates["assigned_staff_id"] = update_data.assigned_staff_id
            updates["assigned_staff_name"] = staff["full_name"]
            
            # Auto-update status to "In Progress" when staff is assigned
            if request_doc.get("status") == "Pending":
                updates["status"] = "In Progress"
                timeline_entry = {
                    "status": "In Progress",
                    "timestamp": now,
                    "note": "Request assigned to staff - Status automatically updated to In Progress",
                    "updated_by": current_user["full_name"]
                }
                await db.recommendation_requests.update_one(
                    {"id": request_id},
                    {"$push": {"timeline": timeline_entry}}
                )
            
            # Notify staff
            await create_notification(
                staff["id"],
                "New Recommendation Assignment",
                f"You have been assigned a recommendation letter request",
                "recommendation_assignment",
                request_id
            )
    
    if update_data.rejection_reason:
        updates["rejection_reason"] = update_data.rejection_reason
        updates["status"] = "Rejected"
        timeline_entry = {
            "status": "Rejected",
            "timestamp": now,
            "note": f"Request rejected: {update_data.rejection_reason}",
            "updated_by": current_user["full_name"]
        }
        await db.recommendation_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.staff_notes:
        updates["staff_notes"] = update_data.staff_notes
    
    if update_data.co_curricular_activities is not None:
        updates["co_curricular_activities"] = update_data.co_curricular_activities
    
    if update_data.administrative_clearance is not None:
        updates["administrative_clearance"] = update_data.administrative_clearance
    
    # Allow students to update their own pending requests
    if current_user["role"] == "student":
        # Update fields students can modify
        if update_data.first_name: updates["first_name"] = update_data.first_name
        if update_data.middle_name is not None: updates["middle_name"] = update_data.middle_name
        if update_data.last_name: updates["last_name"] = update_data.last_name
        if update_data.email: updates["email"] = update_data.email
        if update_data.phone_number: updates["phone_number"] = update_data.phone_number
        if update_data.address: updates["address"] = update_data.address
        if update_data.years_attended: 
            updates["years_attended"] = update_data.years_attended
            # Create string version for backward compatibility
            updates["years_attended_str"] = ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in update_data.years_attended if isinstance(y, dict)])
        if update_data.enrollment_status: updates["enrollment_status"] = update_data.enrollment_status
        if update_data.last_form_class: updates["last_form_class"] = update_data.last_form_class
        if update_data.institution_name: updates["institution_name"] = update_data.institution_name
        if update_data.institution_address: updates["institution_address"] = update_data.institution_address
        if update_data.directed_to: updates["directed_to"] = update_data.directed_to
        if update_data.program_name: updates["program_name"] = update_data.program_name
        if update_data.needed_by_date: updates["needed_by_date"] = update_data.needed_by_date
        if update_data.collection_method: updates["collection_method"] = update_data.collection_method
        if update_data.delivery_address is not None: updates["delivery_address"] = update_data.delivery_address
        if update_data.co_curricular_activities is not None: updates["co_curricular_activities"] = update_data.co_curricular_activities
        if update_data.external_exams is not None: updates["external_exams"] = update_data.external_exams
    
    await db.recommendation_requests.update_one({"id": request_id}, {"$set": updates})
    
    # Notify student of status change (with email)
    effective_new_status = update_data.status or ("Rejected" if update_data.rejection_reason else None)
    if effective_new_status and effective_new_status != old_status:
        student = await db.users.find_one({"id": request_doc["student_id"]}, {"_id": 0})
        if student:
            title = "Recommendation Request Status Updated"
            message = f"Your recommendation letter request has been updated from '{old_status}' to '{effective_new_status}'."
            await create_notification(student["id"], title, message, "recommendation_status_update", request_id)
            
            # Send rich email notification
            updated_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
            email_targets = set()
            student_email = student.get("email", "").strip()
            wolmers_email = updated_doc.get("student_email", "").strip()  # student_email field holds wolmers email
            personal_email = updated_doc.get("email", "").strip()  # email field = personal email for recommendations
            if student_email:
                email_targets.add(student_email)
            if wolmers_email and wolmers_email != student_email:
                email_targets.add(wolmers_email)
            if personal_email and personal_email not in email_targets:
                email_targets.add(personal_email)
            
            app_url = os.environ.get("APP_URL", MICROSOFT_REDIRECT_URI or "")
            html_content = build_status_email_html(
                student_name=student.get("full_name", "Student"),
                request_type="Recommendation Letter",
                old_status=old_status,
                new_status=effective_new_status,
                request_id=request_id,
                rejection_reason=update_data.rejection_reason,
                staff_notes=update_data.staff_notes,
                accent_color="#7b1e2c",
                app_url=app_url
            )
            subject = f"WBS Tracker – Recommendation Letter Request: {effective_new_status}"
            for email in email_targets:
                await send_email_notification(email, subject, html_content)
    
    updated_request = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    normalized_request = normalize_recommendation_data(updated_request)
    return RecommendationRequestResponse(**normalized_request)

@api_router.post("/recommendations/{request_id}/documents")
async def upload_recommendation_document(request_id: str, file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Only staff and admin can upload documents")
    
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Validate file type
    allowed_types = [
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "image/jpeg",
        "image/png",
        "image/gif"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not allowed")
    
    # Save file
    file_id = str(uuid.uuid4())
    file_ext = Path(file.filename).suffix
    file_path = UPLOAD_DIR / f"{file_id}{file_ext}"
    
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    now = datetime.now(timezone.utc).isoformat()
    doc_entry = {
        "id": file_id,
        "filename": file.filename,
        "content_type": file.content_type,
        "path": str(file_path),
        "uploaded_by": current_user["full_name"],
        "uploaded_at": now
    }
    
    await db.recommendation_requests.update_one(
        {"id": request_id},
        {
            "$push": {"documents": doc_entry},
            "$set": {"updated_at": now}
        }
    )
    
    # Add timeline entry
    timeline_entry = {
        "status": request_doc["status"],
        "timestamp": now,
        "note": f"Document uploaded: {file.filename}",
        "updated_by": current_user["full_name"]
    }
    await db.recommendation_requests.update_one(
        {"id": request_id},
        {"$push": {"timeline": timeline_entry}}
    )
    
    # Notify student
    await create_notification(
        request_doc["student_id"],
        "Document Uploaded",
        f"A document has been uploaded to your recommendation letter request",
        "recommendation_document",
        request_id
    )
    
    return {"message": "Document uploaded successfully", "document": doc_entry}

# ==================== NOTIFICATIONS ====================

@api_router.get("/notifications", response_model=List[NotificationResponse])
async def get_notifications(current_user: dict = Depends(get_current_user)):
    notifications = await db.notifications.find(
        {"user_id": current_user["id"]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    return [NotificationResponse(**n) for n in notifications]

@api_router.get("/notifications/unread-count")
async def get_unread_count(current_user: dict = Depends(get_current_user)):
    count = await db.notifications.count_documents({
        "user_id": current_user["id"],
        "read": False
    })
    return {"count": count}

@api_router.patch("/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.notifications.update_one(
        {"id": notification_id, "user_id": current_user["id"]},
        {"$set": {"read": True}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Notification not found")
    return {"message": "Notification marked as read"}

@api_router.patch("/notifications/read-all")
async def mark_all_notifications_read(current_user: dict = Depends(get_current_user)):
    await db.notifications.update_many(
        {"user_id": current_user["id"], "read": False},
        {"$set": {"read": True}}
    )
    return {"message": "All notifications marked as read"}

# ==================== ANALYTICS ====================

@api_router.get("/analytics", response_model=AnalyticsResponse)
async def get_analytics(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view analytics")
    
    # Check and notify about overdue requests
    await check_and_notify_overdue_requests()
    
    now = datetime.now(timezone.utc)
    today_str = now.strftime("%Y-%m-%d")
    
    # Use aggregation pipeline for optimized queries
    pipeline = [
        {
            "$facet": {
                "status_counts": [
                    {"$group": {"_id": "$status", "count": {"$sum": 1}}}
                ],
                "enrollment_counts": [
                    {"$group": {"_id": "$enrollment_status", "count": {"$sum": 1}}}
                ],
                "collection_counts": [
                    {"$group": {"_id": "$collection_method", "count": {"$sum": 1}}}
                ],
                "total": [
                    {"$count": "count"}
                ],
                "staff_workload": [
                    {"$match": {"assigned_staff_id": {"$ne": None}}},
                    {"$group": {"_id": "$assigned_staff_id", "count": {"$sum": 1}}}
                ]
            }
        }
    ]
    
    result = await db.transcript_requests.aggregate(pipeline).to_list(1)
    
    # Calculate overdue requests (needed_by_date < today and status not Completed/Rejected)
    overdue_requests = await db.transcript_requests.find({
        "needed_by_date": {"$lt": today_str},
        "status": {"$nin": ["Completed", "Rejected"]}
    }).to_list(None)
    overdue_count = len(overdue_requests)
    
    # Calculate overdue by days categories
    overdue_by_days = {"1-3 days": 0, "4-7 days": 0, "8-14 days": 0, "15+ days": 0}
    for req in overdue_requests:
        try:
            needed_date = datetime.strptime(req["needed_by_date"], "%Y-%m-%d")
            days_overdue = (now.replace(tzinfo=None) - needed_date).days
            if days_overdue <= 3:
                overdue_by_days["1-3 days"] += 1
            elif days_overdue <= 7:
                overdue_by_days["4-7 days"] += 1
            elif days_overdue <= 14:
                overdue_by_days["8-14 days"] += 1
            else:
                overdue_by_days["15+ days"] += 1
        except:
            pass
    
    overdue_by_days_list = [{"name": k, "value": v, "color": "#ef4444" if "15+" in k else "#f97316" if "8-14" in k else "#eab308" if "4-7" in k else "#fbbf24"} for k, v in overdue_by_days.items() if v > 0]
    
    # Parse aggregation results
    if result:
        data = result[0]
        
        # Get total
        total = data["total"][0]["count"] if data["total"] else 0
        
        # Parse status counts
        status_map = {item["_id"]: item["count"] for item in data["status_counts"]}
        pending = status_map.get("Pending", 0)
        in_progress = status_map.get("In Progress", 0)
        processing = status_map.get("Processing", 0)
        ready = status_map.get("Ready", 0)
        completed = status_map.get("Completed", 0)
        rejected = status_map.get("Rejected", 0)
        
        # Parse enrollment counts
        enrollment_map = {item["_id"]: item["count"] for item in data["enrollment_counts"]}
        requests_by_enrollment = [
            {"name": "Enrolled", "value": enrollment_map.get("enrolled", 0)},
            {"name": "Graduate", "value": enrollment_map.get("graduate", 0)},
            {"name": "Withdrawn", "value": enrollment_map.get("withdrawn", 0)}
        ]
        
        # Parse collection method counts
        collection_map = {item["_id"]: item["count"] for item in data["collection_counts"]}
        requests_by_collection_method = [
            {"name": "Pickup at Bursary", "value": collection_map.get("pickup", 0)},
            {"name": "Emailed to Institution", "value": collection_map.get("emailed", 0)},
            {"name": "Physical Delivery", "value": collection_map.get("delivery", 0)}
        ]
        
        # Parse staff workload - include both transcripts and recommendations
        staff_workload_map = {item["_id"]: item["count"] for item in data["staff_workload"]}
        
        # Get recommendation staff workload
        rec_staff_workload = await db.recommendation_requests.aggregate([
            {"$match": {"assigned_staff_id": {"$ne": None}}},
            {"$group": {"_id": "$assigned_staff_id", "count": {"$sum": 1}}}
        ]).to_list(100)
        
        # Merge transcript and recommendation workloads
        for item in rec_staff_workload:
            staff_id = item["_id"]
            if staff_id in staff_workload_map:
                staff_workload_map[staff_id] += item["count"]
            else:
                staff_workload_map[staff_id] = item["count"]
        
        staff_workload = []
        for staff_id, count in staff_workload_map.items():
            staff = await db.users.find_one({"id": staff_id})
            staff_name = staff["full_name"] if staff else "Unknown"
            staff_workload.append({"name": staff_name, "requests": count})
        
        # Add unassigned count (both transcripts and recommendations)
        unassigned_transcript_count = await db.transcript_requests.count_documents({
            "$or": [{"assigned_staff_id": None}, {"assigned_staff_id": {"$exists": False}}]
        })
        unassigned_rec_count = await db.recommendation_requests.count_documents({
            "$or": [{"assigned_staff_id": None}, {"assigned_staff_id": {"$exists": False}}]
        })
        total_unassigned = unassigned_transcript_count + unassigned_rec_count
        if total_unassigned > 0:
            staff_workload.append({"name": "Unassigned", "requests": total_unassigned})
        
    else:
        total = pending = in_progress = processing = ready = completed = rejected = 0
        requests_by_enrollment = [
            {"name": "Enrolled", "value": 0},
            {"name": "Graduate", "value": 0},
            {"name": "Withdrawn", "value": 0}
        ]
        requests_by_collection_method = [
            {"name": "Pickup at Bursary", "value": 0},
            {"name": "Emailed to Institution", "value": 0},
            {"name": "Physical Delivery", "value": 0}
        ]
        staff_workload = []
    
    # Requests by month (last 6 months) - using aggregation
    requests_by_month = []
    
    for i in range(5, -1, -1):
        month_start = (now.replace(day=1) - timedelta(days=i*30)).replace(day=1)
        month_end = (month_start + timedelta(days=32)).replace(day=1)
        
        # Count transcript requests for this month
        transcript_count = await db.transcript_requests.count_documents({
            "created_at": {
                "$gte": month_start.isoformat(),
                "$lt": month_end.isoformat()
            }
        })
        
        # Count recommendation requests for this month
        recommendation_count = await db.recommendation_requests.count_documents({
            "created_at": {
                "$gte": month_start.isoformat(),
                "$lt": month_end.isoformat()
            }
        })
        
        requests_by_month.append({
            "month": month_start.strftime("%b %Y"),
            "count": transcript_count,
            "transcripts": transcript_count,
            "recommendations": recommendation_count
        })
    
    # Get recommendation letter stats
    total_rec = await db.recommendation_requests.count_documents({})
    pending_rec = await db.recommendation_requests.count_documents({"status": "Pending"})
    completed_rec = await db.recommendation_requests.count_documents({"status": "Completed"})
    in_progress_rec = await db.recommendation_requests.count_documents({"status": "In Progress"})
    rejected_rec = await db.recommendation_requests.count_documents({"status": "Rejected"})
    
    # Recommendations by enrollment status (handle different case variations)
    recommendations_by_enrollment = []
    
    # Get all enrollment statuses from recommendation requests
    rec_enrollment_counts = await db.recommendation_requests.aggregate([
        {"$group": {"_id": "$enrollment_status", "count": {"$sum": 1}}}
    ]).to_list(10)
    
    # Normalize and aggregate enrollment counts
    enrolled_count = 0
    graduate_count = 0
    withdrawn_count = 0
    
    for item in rec_enrollment_counts:
        status = (item["_id"] or "").lower()
        count = item["count"]
        if status in ["enrolled", "currently enrolled"]:
            enrolled_count += count
        elif status in ["graduate", "graduate/alumni", "alumni"]:
            graduate_count += count
        elif status in ["withdrawn"]:
            withdrawn_count += count
    
    recommendations_by_enrollment = [
        {"name": "Enrolled", "value": enrolled_count},
        {"name": "Graduate", "value": graduate_count},
        {"name": "Withdrawn", "value": withdrawn_count}
    ]
    
    # Get overdue recommendation requests
    overdue_rec_count = 0
    overdue_rec_by_days = []
    rec_requests = await db.recommendation_requests.find({
        "status": {"$nin": ["Completed", "Rejected"]},
        "needed_by_date": {"$ne": None, "$ne": ""}
    }, {"_id": 0, "needed_by_date": 1}).to_list(10000)
    
    for req in rec_requests:
        try:
            needed_date = datetime.fromisoformat(req["needed_by_date"].replace('Z', '+00:00')).date()
            if needed_date < now.date():
                overdue_rec_count += 1
                days_overdue = (now.date() - needed_date).days
                if days_overdue <= 7:
                    overdue_rec_by_days.append({"days": "1-7 days", "count": 1})
                elif days_overdue <= 14:
                    overdue_rec_by_days.append({"days": "8-14 days", "count": 1})
                elif days_overdue <= 30:
                    overdue_rec_by_days.append({"days": "15-30 days", "count": 1})
                else:
                    overdue_rec_by_days.append({"days": "30+ days", "count": 1})
        except:
            pass
    
    # Aggregate overdue recommendation by days
    rec_overdue_agg = {}
    for item in overdue_rec_by_days:
        rec_overdue_agg[item["days"]] = rec_overdue_agg.get(item["days"], 0) + 1
    overdue_rec_by_days_list = [{"days": k, "count": v} for k, v in rec_overdue_agg.items()]
    
    # Get recommendation collection method breakdown
    rec_collection_counts = await db.recommendation_requests.aggregate([
        {"$group": {"_id": "$collection_method", "count": {"$sum": 1}}}
    ]).to_list(10)
    rec_collection_map = {item["_id"]: item["count"] for item in rec_collection_counts if item["_id"]}
    recommendations_by_collection_method = [
        {"name": "Pickup at School", "value": rec_collection_map.get("pickup", 0)},
        {"name": "Emailed to Institution", "value": rec_collection_map.get("emailed", 0)},
        {"name": "Physical Delivery", "value": rec_collection_map.get("delivery", 0)}
    ]
    
    return AnalyticsResponse(
        total_requests=total,
        pending_requests=pending,
        in_progress_requests=in_progress,
        processing_requests=processing,
        ready_requests=ready,
        completed_requests=completed,
        rejected_requests=rejected,
        overdue_requests=overdue_count,
        requests_by_month=requests_by_month,
        requests_by_enrollment=requests_by_enrollment,
        recommendations_by_enrollment=recommendations_by_enrollment,
        requests_by_collection_method=requests_by_collection_method,
        staff_workload=staff_workload,
        overdue_by_days=overdue_by_days_list,
        total_recommendation_requests=total_rec,
        pending_recommendation_requests=pending_rec,
        in_progress_recommendation_requests=in_progress_rec,
        completed_recommendation_requests=completed_rec,
        rejected_recommendation_requests=rejected_rec,
        overdue_recommendation_requests=overdue_rec_count,
        recommendations_by_collection_method=recommendations_by_collection_method,
        overdue_transcripts_by_days=overdue_by_days_list,
        overdue_recommendations_by_days=overdue_rec_by_days_list
    )

# ==================== HEALTH CHECK ====================

@api_router.get("/")
async def root():
    return {"message": "WBS Transcript and Recommendation Tracker API", "status": "running"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy"}

# ==================== EXPORT/REPORTING ENDPOINTS ====================

def format_date_for_export(date_str):
    """Format date string for export"""
    try:
        if date_str:
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return dt.strftime("%Y-%m-%d %H:%M")
    except:
        pass
    return date_str or ""

def format_years_for_export(years_data):
    """Format years attended/academic years for export"""
    if isinstance(years_data, list):
        return ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in years_data])
    return str(years_data) if years_data else ""

@api_router.get("/export/transcripts/{format_type}")
async def export_transcript_requests(format_type: str, status: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """Export transcript requests as DOCX, PDF, or XLSX"""
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    # Build query
    query = {}
    if current_user["role"] == "staff":
        query["assigned_staff_id"] = current_user["id"]
    if status and status != "all":
        query["status"] = status
    
    requests = await db.transcript_requests.find(query, {"_id": 0}).sort("created_at", -1).to_list(10000)
    
    if format_type == "xlsx":
        return generate_transcript_xlsx(requests)
    elif format_type == "pdf":
        return generate_transcript_pdf(requests)
    elif format_type == "docx":
        return generate_transcript_docx(requests)
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use xlsx, pdf, or docx")

@api_router.get("/export/recommendations/{format_type}")
async def export_recommendation_requests(format_type: str, status: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """Export recommendation requests as DOCX, PDF, or XLSX"""
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    # Build query
    query = {}
    if current_user["role"] == "staff":
        query["assigned_staff_id"] = current_user["id"]
    if status and status != "all":
        query["status"] = status
    
    requests = await db.recommendation_requests.find(query, {"_id": 0}).sort("created_at", -1).to_list(10000)
    
    if format_type == "xlsx":
        return generate_recommendation_xlsx(requests)
    elif format_type == "pdf":
        return generate_recommendation_pdf(requests)
    elif format_type == "docx":
        return generate_recommendation_docx(requests)
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use xlsx, pdf, or docx")

def generate_transcript_xlsx(requests):
    """Generate XLSX file for transcript requests"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Transcript Requests"
    
    # Headers
    headers = ["ID", "Student Name", "Email", "School ID", "Status", "Academic Years", 
               "Collection Method", "Institution", "Needed By", "Assigned Staff", "Created At"]
    
    # Style headers
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="800000", end_color="800000", fill_type="solid")
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for row_num, req in enumerate(requests, 2):
        ws.cell(row=row_num, column=1, value=req.get("id", "")[:8])
        ws.cell(row=row_num, column=2, value=req.get("student_name", ""))
        ws.cell(row=row_num, column=3, value=req.get("student_email", ""))
        ws.cell(row=row_num, column=4, value=req.get("school_id", ""))
        ws.cell(row=row_num, column=5, value=req.get("status", ""))
        ws.cell(row=row_num, column=6, value=format_years_for_export(req.get("academic_years", req.get("academic_year", ""))))
        ws.cell(row=row_num, column=7, value=req.get("collection_method", ""))
        ws.cell(row=row_num, column=8, value=req.get("institution_name", ""))
        ws.cell(row=row_num, column=9, value=format_date_for_export(req.get("needed_by_date", "")))
        ws.cell(row=row_num, column=10, value=req.get("assigned_staff_name", "Unassigned"))
        ws.cell(row=row_num, column=11, value=format_date_for_export(req.get("created_at", "")))
    
    # Adjust column widths
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 40)
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=transcript_requests_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

def generate_transcript_pdf(requests):
    """Generate PDF file for transcript requests"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=30, bottomMargin=30)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=1, spaceAfter=20)
    elements.append(Paragraph("Transcript Requests Report", title_style))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Table data
    data = [["ID", "Student", "Status", "Academic Years", "Collection", "Institution", "Needed By", "Staff"]]
    
    for req in requests:
        data.append([
            req.get("id", "")[:8],
            req.get("student_name", ""),
            req.get("status", ""),
            format_years_for_export(req.get("academic_years", req.get("academic_year", "")))[:20],
            req.get("collection_method", ""),
            (req.get("institution_name", "") or "")[:20],
            format_date_for_export(req.get("needed_by_date", ""))[:10],
            (req.get("assigned_staff_name", "") or "Unassigned")[:15]
        ])
    
    # Create table
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.5, 0, 0)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=transcript_requests_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

def generate_transcript_docx(requests):
    """Generate DOCX file for transcript requests"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Transcript Requests Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Total Requests: {len(requests)}")
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["Student", "Status", "Academic Years", "Collection", "Institution", "Needed By", "Staff"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for req in requests:
        row_cells = table.add_row().cells
        row_cells[0].text = req.get("student_name", "")
        row_cells[1].text = req.get("status", "")
        row_cells[2].text = format_years_for_export(req.get("academic_years", req.get("academic_year", "")))
        row_cells[3].text = req.get("collection_method", "")
        row_cells[4].text = req.get("institution_name", "") or ""
        row_cells[5].text = format_date_for_export(req.get("needed_by_date", ""))[:10]
        row_cells[6].text = req.get("assigned_staff_name", "") or "Unassigned"
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=transcript_requests_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

def generate_recommendation_xlsx(requests):
    """Generate XLSX file for recommendation requests"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Recommendation Requests"
    
    # Headers
    headers = ["ID", "Student Name", "Email", "Status", "Years Attended", "Form Class",
               "Institution", "Program", "Collection Method", "Needed By", "Assigned Staff", "Created At"]
    
    # Style headers
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="DAA520", end_color="DAA520", fill_type="solid")
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for row_num, req in enumerate(requests, 2):
        ws.cell(row=row_num, column=1, value=req.get("id", "")[:8])
        ws.cell(row=row_num, column=2, value=req.get("student_name", ""))
        ws.cell(row=row_num, column=3, value=req.get("student_email", ""))
        ws.cell(row=row_num, column=4, value=req.get("status", ""))
        ws.cell(row=row_num, column=5, value=format_years_for_export(req.get("years_attended", req.get("years_attended_str", ""))))
        ws.cell(row=row_num, column=6, value=req.get("last_form_class", ""))
        ws.cell(row=row_num, column=7, value=req.get("institution_name", ""))
        ws.cell(row=row_num, column=8, value=req.get("program_name", ""))
        ws.cell(row=row_num, column=9, value=req.get("collection_method", ""))
        ws.cell(row=row_num, column=10, value=format_date_for_export(req.get("needed_by_date", "")))
        ws.cell(row=row_num, column=11, value=req.get("assigned_staff_name", "Unassigned"))
        ws.cell(row=row_num, column=12, value=format_date_for_export(req.get("created_at", "")))
    
    # Adjust column widths
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 40)
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=recommendation_requests_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

def generate_recommendation_pdf(requests):
    """Generate PDF file for recommendation requests"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=30, bottomMargin=30)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=1, spaceAfter=20)
    elements.append(Paragraph("Recommendation Letter Requests Report", title_style))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Table data
    data = [["ID", "Student", "Status", "Years", "Institution", "Program", "Collection", "Needed By", "Staff"]]
    
    for req in requests:
        data.append([
            req.get("id", "")[:8],
            req.get("student_name", ""),
            req.get("status", ""),
            format_years_for_export(req.get("years_attended", req.get("years_attended_str", "")))[:15],
            (req.get("institution_name", "") or "")[:18],
            (req.get("program_name", "") or "")[:18],
            req.get("collection_method", ""),
            format_date_for_export(req.get("needed_by_date", ""))[:10],
            (req.get("assigned_staff_name", "") or "Unassigned")[:12]
        ])
    
    # Create table
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.85, 0.65, 0.13)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightyellow),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 7),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=recommendation_requests_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

def generate_recommendation_docx(requests):
    """Generate DOCX file for recommendation requests"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Recommendation Letter Requests Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Total Requests: {len(requests)}")
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["Student", "Status", "Years", "Form Class", "Institution", "Program", "Needed By", "Staff"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for req in requests:
        row_cells = table.add_row().cells
        row_cells[0].text = req.get("student_name", "")
        row_cells[1].text = req.get("status", "")
        row_cells[2].text = format_years_for_export(req.get("years_attended", req.get("years_attended_str", "")))
        row_cells[3].text = req.get("last_form_class", "")
        row_cells[4].text = req.get("institution_name", "") or ""
        row_cells[5].text = req.get("program_name", "") or ""
        row_cells[6].text = format_date_for_export(req.get("needed_by_date", ""))[:10]
        row_cells[7].text = req.get("assigned_staff_name", "") or "Unassigned"
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=recommendation_requests_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

# ==================== ADMIN DATA MANAGEMENT ====================

class DataClearResponse(BaseModel):
    success: bool
    message: str
    deleted_counts: dict

@api_router.get("/admin/export-all-data/pdf")
async def export_all_data_pdf(current_user: dict = Depends(get_current_user)):
    """Export all data (users, transcripts, recommendations) to PDF before clearing"""
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Fetch all data
    users = await db.users.find({"role": {"$ne": "admin"}}, {"_id": 0, "password_hash": 0}).to_list(1000)
    transcripts = await db.transcript_requests.find({}, {"_id": 0}).to_list(1000)
    recommendations = await db.recommendation_requests.find({}, {"_id": 0}).to_list(1000)
    notifications = await db.notifications.find({}, {"_id": 0}).to_list(1000)
    
    # Generate PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=0.5*inch, bottomMargin=0.5*inch)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#800000'),
        spaceAfter=20,
        alignment=1
    )
    elements.append(Paragraph("WBS Transcript and Recommendation Tracker - Complete Data Export", title_style))
    elements.append(Paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Summary
    summary_data = [
        ["Category", "Count"],
        ["Users (non-admin)", str(len(users))],
        ["Transcript Requests", str(len(transcripts))],
        ["Recommendation Requests", str(len(recommendations))],
        ["Notifications", str(len(notifications))],
    ]
    summary_table = Table(summary_data, colWidths=[3*inch, 1.5*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#800000')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(Paragraph("Data Summary", styles['Heading2']))
    elements.append(summary_table)
    elements.append(Spacer(1, 30))
    
    # Users Section
    if users:
        elements.append(Paragraph("Users (Non-Admin)", styles['Heading2']))
        user_data = [["Name", "Email", "Role", "Created At"]]
        for user in users:
            user_data.append([
                user.get("full_name", ""),
                user.get("email", ""),
                user.get("role", ""),
                str(user.get("created_at", ""))[:19]
            ])
        user_table = Table(user_data, colWidths=[2.5*inch, 3*inch, 1*inch, 2*inch])
        user_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#800000')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ]))
        elements.append(user_table)
        elements.append(Spacer(1, 20))
    
    # Transcript Requests Section
    if transcripts:
        elements.append(Paragraph("Transcript Requests", styles['Heading2']))
        transcript_data = [["Student", "Status", "Academic Years", "Collection", "Created At"]]
        for req in transcripts:
            academic_years = req.get("academic_years", req.get("academic_year", ""))
            if isinstance(academic_years, list):
                academic_years = ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in academic_years])
            transcript_data.append([
                req.get("student_name", ""),
                req.get("status", ""),
                str(academic_years)[:30],
                req.get("collection_method", ""),
                str(req.get("created_at", ""))[:10]
            ])
        transcript_table = Table(transcript_data, colWidths=[2*inch, 1.2*inch, 2*inch, 1.3*inch, 1.5*inch])
        transcript_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#800000')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ]))
        elements.append(transcript_table)
        elements.append(Spacer(1, 20))
    
    # Recommendation Requests Section
    if recommendations:
        elements.append(Paragraph("Recommendation Requests", styles['Heading2']))
        rec_data = [["Student", "Institution", "Program", "Status", "Created At"]]
        for req in recommendations:
            rec_data.append([
                req.get("student_name", ""),
                str(req.get("institution_name", ""))[:25],
                str(req.get("program_name", ""))[:20],
                req.get("status", ""),
                str(req.get("created_at", ""))[:10]
            ])
        rec_table = Table(rec_data, colWidths=[2*inch, 2.5*inch, 1.8*inch, 1.2*inch, 1.2*inch])
        rec_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#800000')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ]))
        elements.append(rec_table)
        elements.append(Spacer(1, 20))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=wbs_complete_data_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"}
    )

@api_router.delete("/admin/clear-all-data")
async def clear_all_data(current_user: dict = Depends(get_current_user)):
    """Clear all data from the database except the admin account"""
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        # Delete all users except admin
        users_result = await db.users.delete_many({"email": {"$ne": "admin@wolmers.org"}})
        
        # Delete all transcript requests
        transcripts_result = await db.transcript_requests.delete_many({})
        
        # Delete all recommendation requests
        recommendations_result = await db.recommendation_requests.delete_many({})
        
        # Delete all notifications
        notifications_result = await db.notifications.delete_many({})
        
        # Delete all password reset tokens
        password_resets_result = await db.password_resets.delete_many({})
        
        deleted_counts = {
            "users": users_result.deleted_count,
            "transcript_requests": transcripts_result.deleted_count,
            "recommendation_requests": recommendations_result.deleted_count,
            "notifications": notifications_result.deleted_count,
            "password_resets": password_resets_result.deleted_count
        }
        
        total_deleted = sum(deleted_counts.values())
        
        logger.info(f"Admin {current_user['email']} cleared all data: {deleted_counts}")
        
        return DataClearResponse(
            success=True,
            message=f"Successfully cleared {total_deleted} records from the database",
            deleted_counts=deleted_counts
        )
    except Exception as e:
        logger.error(f"Error clearing data: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to clear data: {str(e)}")

@api_router.get("/admin/data-summary")
async def get_data_summary(current_user: dict = Depends(get_current_user)):
    """Get summary of all data in the database"""
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    users_count = await db.users.count_documents({"email": {"$ne": "admin@wolmers.org"}})
    transcripts_count = await db.transcript_requests.count_documents({})
    recommendations_count = await db.recommendation_requests.count_documents({})
    notifications_count = await db.notifications.count_documents({})
    
    return {
        "users": users_count,
        "transcript_requests": transcripts_count,
        "recommendation_requests": recommendations_count,
        "notifications": notifications_count,
        "total": users_count + transcripts_count + recommendations_count + notifications_count
    }


# ==================== JSON EXPORT / IMPORT ====================

class ImportResponse(BaseModel):
    success: bool
    message: str
    imported_counts: dict


@api_router.get("/admin/export-all-data/json")
async def export_all_data_json(current_user: dict = Depends(get_current_user)):
    """Export all app data as a JSON file for backup / migration."""
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    try:
        # Fetch all collections (strip MongoDB _id)
        users = await db.users.find({}, {"_id": 0}).to_list(None)
        transcripts = await db.transcript_requests.find({}, {"_id": 0}).to_list(None)
        recommendations = await db.recommendation_requests.find({}, {"_id": 0}).to_list(None)
        notifications = await db.notifications.find({}, {"_id": 0}).to_list(None)

        export_payload = {
            "export_version": "1.0",
            "exported_at": datetime.now(timezone.utc).isoformat(),
            "exported_by": current_user["email"],
            "app_name": "WBS Transcript & Recommendation Tracker",
            "counts": {
                "users": len(users),
                "transcript_requests": len(transcripts),
                "recommendation_requests": len(recommendations),
                "notifications": len(notifications),
            },
            "data": {
                "users": users,
                "transcript_requests": transcripts,
                "recommendation_requests": recommendations,
                "notifications": notifications,
            }
        }

        import json
        json_bytes = json.dumps(export_payload, indent=2, default=str).encode("utf-8")
        filename = f"wbs_data_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        return StreamingResponse(
            io.BytesIO(json_bytes),
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        logger.error(f"JSON export error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to export data: {str(e)}")


@api_router.post("/admin/import-data/json", response_model=ImportResponse)
async def import_all_data_json(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Import app data from a previously exported JSON file."""
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    if not file.filename.endswith(".json"):
        raise HTTPException(status_code=400, detail="Only .json files are accepted")

    try:
        import json
        raw = await file.read()
        payload = json.loads(raw)

        if "data" not in payload:
            raise HTTPException(status_code=400, detail="Invalid export file: missing 'data' key")

        data = payload["data"]
        imported_counts: dict = {}

        # Users – upsert by id to avoid duplicates
        if "users" in data and data["users"]:
            for u in data["users"]:
                # Never overwrite the current admin account
                if u.get("email") == current_user["email"]:
                    continue
                await db.users.update_one({"id": u["id"]}, {"$set": u}, upsert=True)
            imported_counts["users"] = len(data["users"])

        # Transcript requests
        if "transcript_requests" in data and data["transcript_requests"]:
            for tr in data["transcript_requests"]:
                await db.transcript_requests.update_one({"id": tr["id"]}, {"$set": tr}, upsert=True)
            imported_counts["transcript_requests"] = len(data["transcript_requests"])

        # Recommendation requests
        if "recommendation_requests" in data and data["recommendation_requests"]:
            for rr in data["recommendation_requests"]:
                await db.recommendation_requests.update_one({"id": rr["id"]}, {"$set": rr}, upsert=True)
            imported_counts["recommendation_requests"] = len(data["recommendation_requests"])

        # Notifications
        if "notifications" in data and data["notifications"]:
            for n in data["notifications"]:
                await db.notifications.update_one({"id": n["id"]}, {"$set": n}, upsert=True)
            imported_counts["notifications"] = len(data["notifications"])

        total = sum(imported_counts.values())
        logger.info(f"Admin {current_user['email']} imported data: {imported_counts}")

        return ImportResponse(
            success=True,
            message=f"Successfully imported {total} records",
            imported_counts=imported_counts
        )
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON file")
    except Exception as e:
        logger.error(f"JSON import error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to import data: {str(e)}")


# ==================== SEED DEFAULT ADMIN ====================

@app.on_event("startup")
async def seed_default_admin():
    admin_email = "admin@wolmers.org"
    existing = await db.users.find_one({"email": admin_email}, {"_id": 0})
    
    if not existing:
        admin_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        
        admin_doc = {
            "id": admin_id,
            "email": admin_email,
            "full_name": "System Administrator",
            "password_hash": hash_password("Admin123!"),
            "role": "admin",
            "created_at": now,
            "updated_at": now
        }
        
        await db.users.insert_one(admin_doc)
        logger.info("Default admin account created: admin@wolmers.org / Admin123!")

# ============================================================
# SINGLE RECORD EXPORT ENDPOINTS
# ============================================================

WOLMERS_LOGO_URL = "https://customer-assets.emergentagent.com/job_13afcd2c-9b31-4868-9eb9-1450f0dbe963/artifacts/iuukr0xo_Wolmer%27s_Schools.png"

def format_external_exams(exams: list) -> str:
    if not exams:
        return "N/A"
    parts = []
    for e in exams:
        if e.get("exam") == "Not Applicable":
            return "Not Applicable"
        name = e.get("name", e.get("exam", ""))
        year = e.get("year", "")
        if year:
            parts.append(f"{name} ({year})")
        else:
            parts.append(name)
    return ", ".join(parts)

def format_clearance_for_doc(ac: dict) -> dict:
    if not ac:
        return {"no_fees": False, "no_admin": False, "amount": "—", "receipt": "—", "date": "—", "updated_by": "—"}
    return {
        "no_fees": ac.get("no_fees_outstanding", False),
        "no_admin": ac.get("no_admin_obligations", False),
        "amount": f"${ac['amount_paid']:.2f}" if ac.get("amount_paid") is not None else "—",
        "receipt": ac.get("receipt_number") or "—",
        "date": format_date_for_export(ac.get("payment_date", "")) or "—",
        "updated_by": ac.get("updated_by") or "—",
    }

def build_transcript_pdf(request: dict) -> io.BytesIO:
    """Generate a formatted single-record transcript PDF"""
    buffer = io.BytesIO()
    # Usable width = A4(595) - 50*2 margins = 495
    W = 495
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    story = []
    styles = getSampleStyleSheet()

    # Colors
    maroon       = HexColor('#7b1e2c')
    light_maroon = HexColor('#f9f0f2')
    dark_text    = HexColor('#1c1917')
    mid_grey     = HexColor('#78716c')
    light_border = HexColor('#e7e5e4')
    gold         = HexColor('#b8860b')
    red_bg       = HexColor('#fef2f2')
    red_border   = HexColor('#fca5a5')
    blue_bg      = HexColor('#f0f9ff')
    green_bg     = HexColor('#f0fdf4')

    title_style   = ParagraphStyle('Title',    parent=styles['Normal'], fontSize=16, textColor=maroon,   spaceAfter=4,  alignment=1, fontName='Helvetica-Bold')
    subtitle_style= ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=10, textColor=mid_grey, spaceAfter=2,  alignment=1)
    section_style = ParagraphStyle('Section',  parent=styles['Normal'], fontSize=11, textColor=maroon,   fontName='Helvetica-Bold', spaceBefore=10, spaceAfter=4)
    label_style   = ParagraphStyle('Label',    parent=styles['Normal'], fontSize=9,  textColor=mid_grey, leading=12, wordWrap='LTR')
    value_style   = ParagraphStyle('Value',    parent=styles['Normal'], fontSize=9,  textColor=dark_text, fontName='Helvetica-Bold', leading=12, wordWrap='LTR')
    small_style   = ParagraphStyle('Small',    parent=styles['Normal'], fontSize=8,  textColor=mid_grey, leading=11,  wordWrap='LTR')
    check_style   = ParagraphStyle('Check',    parent=styles['Normal'], fontSize=9,  textColor=dark_text, leading=12)
    note_style    = ParagraphStyle('Note',     parent=styles['Normal'], fontSize=9,  textColor=dark_text, leading=13, wordWrap='LTR')

    # Column widths for 4-col info tables — MUST sum to W
    LW = 90   # label
    VW = 155  # value
    # total for 2 pairs: 90+155+90+155 = 490 (leave 5 breathing room)
    COL4 = [LW, VW, LW, VW]

    # ---- HEADER ----
    story.append(Paragraph("New Transcript Request Form", title_style))
    story.append(Spacer(1, 6))

    try:
        import urllib.request as _ur
        logo_data = io.BytesIO(_ur.urlopen(WOLMERS_LOGO_URL, timeout=5).read())
        logo_img = RLImage(logo_data, width=50, height=50)
    except Exception:
        logo_img = None

    school_para = Paragraph(
        "<b>WOLMER'S BOYS' SCHOOL</b><br/>"
        "National Heroes Circle, Kingston 4<br/>"
        "Bursary: 876 922 4055 / 876 948 4807 &nbsp;|&nbsp; WhatsApp: 876 313 0915<br/>"
        "Email: wbs.bursary@wolmers.org / iyona.forbes@wolmers.org<br/>"
        "Guidance: 876 922 8254 &nbsp;|&nbsp; wbs.guidance@wolmers.org",
        ParagraphStyle('School', parent=styles['Normal'], fontSize=9, textColor=dark_text, leading=13, alignment=1, wordWrap='LTR')
    )

    if logo_img:
        header_table = Table([[logo_img], [school_para]], colWidths=[W])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('BOTTOMPADDING', (0,0), (0,0), 4),
            ('BOTTOMPADDING', (0,1), (0,1), 8),
            ('LINEBELOW', (0,1), (-1,1), 0.5, light_border),
        ]))
    else:
        header_table = Table([[school_para]], colWidths=[W])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('LINEBELOW', (0,0), (-1,0), 0.5, light_border),
        ]))
    story.append(header_table)
    story.append(Spacer(1, 10))

    status = request.get('status', 'Pending')
    story.append(Paragraph(
        f"Status: <b>{status}</b>  |  Submitted: <b>{format_date_for_export(request.get('created_at',''))}</b>"
        f"  |  Needed By: <b>{format_date_for_export(request.get('needed_by_date',''))}</b>",
        subtitle_style))
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=0.5, color=light_border))
    story.append(Spacer(1, 6))

    # ---- REJECTION BANNER ----
    if status == 'Rejected' and request.get('rejection_reason'):
        rej_box = Table([[
            Paragraph(f"<b>REJECTED</b> — {request['rejection_reason']}", ParagraphStyle(
                'RejNote', parent=styles['Normal'], fontSize=9, textColor=HexColor('#b91c1c'),
                leading=13, wordWrap='LTR'))
        ]], colWidths=[W])
        rej_box.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), red_bg),
            ('LINEAFTER',  (0,0), (0,-1), 3, HexColor('#ef4444')),
            ('LINEBEFORE', (0,0), (0,-1), 3, HexColor('#ef4444')),
            ('TOPPADDING', (0,0), (-1,-1), 8), ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 12), ('RIGHTPADDING', (0,0), (-1,-1), 12),
        ]))
        story.append(rej_box)
        story.append(Spacer(1, 8))

    def section_header(text):
        story.append(Paragraph(text, section_style))
        story.append(HRFlowable(width="100%", thickness=0.3, color=light_border))
        story.append(Spacer(1, 4))

    def make_info_table(rows):
        """Rows: list of dicts with keys 'type', 'l1','v1','l2','v2' or 'label','value' for full-width."""
        flat = []
        spans = []
        for r in rows:
            if r.get('full'):
                flat.append([
                    Paragraph(r['label'], label_style),
                    Paragraph(r['value'] or '—', value_style),
                ])
                spans.append(True)
            else:
                flat.append([
                    Paragraph(r['l1'], label_style), Paragraph(r.get('v1') or '—', value_style),
                    Paragraph(r.get('l2',''), label_style), Paragraph(r.get('v2') or '—', value_style),
                ])
                spans.append(False)

        # Build tables: group consecutive rows of same type
        # Simpler: alternate 2-col and 4-col sub-tables
        tables = []
        i = 0
        while i < len(flat):
            if spans[i]:
                t = Table([flat[i]], colWidths=[LW, W-LW])
            else:
                t = Table([flat[i]], colWidths=COL4)
            t.setStyle(TableStyle([
                ('ROWBACKGROUNDS', (0,0), (-1,-1), [HexColor('#fafaf9'), HexColor('#f5f5f4')]),
                ('GRID', (0,0), (-1,-1), 0.3, light_border),
                ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ]))
            tables.append(t)
            i += 1
        return tables

    def tc(l1, v1, l2="", v2=""):
        return {'l1': l1, 'v1': v1, 'l2': l2, 'v2': v2}

    def fc(label, value):
        return {'full': True, 'label': label, 'value': value}

    def add_info_block(rows):
        for t in make_info_table(rows):
            story.append(t)
        story.append(Spacer(1, 8))

    # ---- PERSONAL INFO ----
    section_header("Personal Information")
    years_list = request.get('academic_years', [])
    years_str = ", ".join([f"{y.get('from_year')}-{y.get('to_year')}" for y in years_list]) if years_list else "—"
    add_info_block([
        tc("First Name",       request.get('first_name'),              "Last Name",       request.get('last_name')),
        tc("Middle Name",      request.get('middle_name') or "—",      "Date of Birth",   request.get('date_of_birth')),
        tc("School ID",        request.get('school_id') or "—",        "Enrollment",      (request.get('enrollment_status') or "").replace('graduate','Graduate/Alumni').title()),
        tc("Academic Years",   years_str,                               "Last Form Class", request.get('last_form_class')),
    ])

    # ---- CONTACT INFO ----
    section_header("Contact Information")
    add_info_block([
        tc("Wolmer's Email", request.get('wolmers_email') or request.get('student_email'), "Personal Email", request.get('personal_email')),
        tc("Phone Number",   request.get('phone_number'),  "School ID", request.get('school_id') or "—"),
    ])

    # ---- EXTERNAL EXAMS ----
    section_header("External Examinations")
    story.append(Paragraph(format_external_exams(request.get('external_exams', [])), value_style))
    story.append(Spacer(1, 8))

    # ---- REQUEST DETAILS ----
    section_header("Request Details")
    coll_map = {'pickup': 'Pickup at Bursary', 'emailed': 'Email to Institution', 'delivery': 'Courier Delivery (DHL)'}
    req_rows = [
        tc("Reason", (request.get('reason') or '').replace('_',' ').title(), "Number of Copies", str(request.get('number_of_copies') or 1)),
        tc("Collection", coll_map.get(request.get('collection_method',''), request.get('collection_method','')), "Received Before?", request.get('received_transcript_before') or "—"),
    ]
    if request.get('other_reason'):
        req_rows.append(fc("Other Reason", request['other_reason']))
    if request.get('delivery_address'):
        req_rows.append(fc("Delivery Address", request['delivery_address']))
    add_info_block(req_rows)

    # ---- INSTITUTION ----
    section_header("Destination Institution")
    add_info_block([
        tc("Institution Name", request.get('institution_name'), "Phone", request.get('institution_phone')),
        tc("Institution Email", request.get('institution_email'), "", ""),
        fc("Address", request.get('institution_address')),
    ])

    # ---- ADMINISTRATIVE CLEARANCE ----
    story.append(HRFlowable(width="100%", thickness=1, color=maroon))
    story.append(Spacer(1, 4))
    story.append(Paragraph("Administrative Clearance", ParagraphStyle('ClearHeader', parent=styles['Normal'],
        fontSize=12, textColor=maroon, fontName='Helvetica-Bold', spaceBefore=4, spaceAfter=6)))

    ac = format_clearance_for_doc(request.get('administrative_clearance'))

    def _draw_checkbox(checked: bool, color):
        from reportlab.graphics.shapes import Drawing, Rect, Line
        sz = 8
        d = Drawing(sz + 4, sz + 4)
        d.add(Rect(2, 2, sz, sz, strokeColor=color, strokeWidth=0.8, fillColor=HexColor('#ffffff')))
        if checked:
            d.add(Line(3, 6, 5.5, 3.5, strokeColor=color, strokeWidth=1.2))
            d.add(Line(5.5, 3.5, 10, 9, strokeColor=color, strokeWidth=1.2))
        return d

    def checkbox_row(label: str, checked: bool) -> list:
        return [_draw_checkbox(checked, maroon), Paragraph(f"<b>{label}</b>", check_style)]

    ct_checkboxes = Table([
        checkbox_row("No fees outstanding", ac['no_fees']),
        checkbox_row("No outstanding administrative obligations", ac['no_admin']),
    ], colWidths=[16, W-16])
    ct_checkboxes.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), HexColor('#f9f0f2')),
        ('GRID', (0,0), (-1,-1), 0, HexColor('#f9f0f2')),
        ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8), ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(ct_checkboxes)

    payment_label = ParagraphStyle('PayLabel', parent=styles['Normal'], fontSize=9, textColor=mid_grey, wordWrap='LTR')
    payment_val   = ParagraphStyle('PayVal',   parent=styles['Normal'], fontSize=9, textColor=dark_text, fontName='Helvetica-Bold', wordWrap='LTR')
    pay_rows = [[
        Paragraph("Amount Paid:", payment_label), Paragraph(ac['amount'], payment_val),
        Paragraph("Receipt #:", payment_label),   Paragraph(ac['receipt'], payment_val),
        Paragraph("Date:", payment_label),        Paragraph(ac['date'], payment_val),
    ]]
    if ac['updated_by'] != '—':
        pay_rows.append([Paragraph("Processed by:", payment_label), Paragraph(ac['updated_by'], payment_val), '', '', '', ''])
    pt = Table(pay_rows, colWidths=[65, 90, 55, 110, 40, W-360])
    pt.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), HexColor('#f9f0f2')),
        ('GRID', (0,0), (-1,-1), 0, HexColor('#f9f0f2')),
        ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8), ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('SPAN', (1,1), (-1,1)) if ac['updated_by'] != '—' else ('SPAN', (0,0), (0,0)),
    ]))
    story.append(pt)
    story.append(Spacer(1, 10))

    # ---- STAFF NOTES ----
    if request.get('staff_notes'):
        notes_box = Table([[
            Paragraph(f"<b>Staff Notes:</b> {request['staff_notes']}", ParagraphStyle(
                'StaffNote', parent=styles['Normal'], fontSize=9, textColor=HexColor('#0369a1'),
                leading=13, wordWrap='LTR'))
        ]], colWidths=[W])
        notes_box.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), blue_bg),
            ('LINEBEFORE', (0,0), (0,-1), 3, HexColor('#0284c7')),
            ('TOPPADDING', (0,0), (-1,-1), 7), ('BOTTOMPADDING', (0,0), (-1,-1), 7),
            ('LEFTPADDING', (0,0), (-1,-1), 12), ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(notes_box)
        story.append(Spacer(1, 8))

    # ---- STATUS HISTORY ----
    timeline = request.get('timeline', [])
    if timeline:
        story.append(HRFlowable(width="100%", thickness=0.5, color=light_border))
        story.append(Spacer(1, 4))
        story.append(Paragraph("Status History", ParagraphStyle('TLHeader', parent=styles['Normal'],
            fontSize=11, textColor=maroon, fontName='Helvetica-Bold', spaceBefore=4, spaceAfter=4)))
        tl_data = [[
            Paragraph("<b>Status</b>", label_style),
            Paragraph("<b>Date</b>",   label_style),
            Paragraph("<b>Updated By</b>", label_style),
            Paragraph("<b>Note</b>",   label_style),
        ]]
        for entry in timeline:
            tl_data.append([
                Paragraph(entry.get('status',''), value_style),
                Paragraph(format_date_for_export(entry.get('timestamp','')), small_style),
                Paragraph(entry.get('updated_by','System'), small_style),
                Paragraph(entry.get('note','') or '—', small_style),
            ])
        tl_t = Table(tl_data, colWidths=[90, 80, 120, W-290])
        tl_t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), maroon),
            ('TEXTCOLOR', (0,0), (-1,0), HexColor('#ffffff')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [HexColor('#fafaf9'), HexColor('#f5f5f4')]),
            ('GRID', (0,0), (-1,-1), 0.3, light_border),
            ('TOPPADDING', (0,0), (-1,-1), 4), ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        story.append(tl_t)

    doc.build(story)
    buffer.seek(0)
    return buffer


def build_transcript_docx(request: dict) -> io.BytesIO:
    """Generate a formatted single-record transcript DOCX"""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import urllib.request

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    maroon_rgb = RGBColor(0x7b, 0x1e, 0x2c)
    grey_rgb   = RGBColor(0x78, 0x71, 0x6c)
    dark_rgb   = RGBColor(0x1c, 0x19, 0x17)
    red_rgb    = RGBColor(0xb9, 0x1c, 0x1c)
    blue_rgb   = RGBColor(0x03, 0x69, 0xa1)

    # ---- HEADER ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("New Transcript Request Form")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = maroon_rgb

    # Logo centred on its own line
    try:
        logo_bytes = io.BytesIO(urllib.request.urlopen(WOLMERS_LOGO_URL, timeout=5).read())
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_bytes, width=Inches(0.75))
    except Exception:
        pass

    # School info centred below logo
    school_lines = [
        ("WOLMER'S BOYS' SCHOOL", True,  Pt(12), maroon_rgb),
        ("National Heroes Circle, Kingston 4", False, Pt(9), dark_rgb),
        ("Bursary: 876 922 4055 / 876 948 4807  |  WhatsApp: 876 313 0915", False, Pt(9), dark_rgb),
        ("Email: wbs.bursary@wolmers.org / iyona.forbes@wolmers.org", False, Pt(9), dark_rgb),
        ("Guidance Department: 876 922 8254  |  wbs.guidance@wolmers.org", False, Pt(9), grey_rgb),
    ]
    for text, bold, size, color in school_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = size
        r.font.color.rgb = color

    # Status line
    status_para = doc.add_paragraph()
    status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = request.get('status', 'Pending')
    sr = status_para.add_run(
        f"Status: {st}  |  Submitted: {format_date_for_export(request.get('created_at',''))}"
        f"  |  Needed By: {format_date_for_export(request.get('needed_by_date',''))}"
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = grey_rgb

    # Rejection banner
    if st == 'Rejected' and request.get('rejection_reason'):
        rej_p = doc.add_paragraph()
        rej_r = rej_p.add_run(f"REJECTED — {request['rejection_reason']}")
        rej_r.bold = True
        rej_r.font.size = Pt(10)
        rej_r.font.color.rgb = red_rgb

    doc.add_paragraph()

    def add_section_heading(text, color=maroon_rgb):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = color

    def add_field_table(rows_data):
        """Each row: (label, value, label, value) or (label, value) for full-width."""
        tbl = doc.add_table(rows=len(rows_data), cols=4)
        tbl.style = 'Table Grid'
        for i, row_data in enumerate(rows_data):
            row = tbl.rows[i]
            if len(row_data) >= 4:
                for j, (cell_text, is_bold) in enumerate(zip(row_data, [False, True, False, True])):
                    cell = row.cells[j]
                    p = cell.paragraphs[0]
                    r = p.add_run(str(cell_text) if cell_text else "—")
                    r.bold = is_bold
                    r.font.size = Pt(9)
                    r.font.color.rgb = dark_rgb if is_bold else grey_rgb
            else:
                # Full width row — merge cols 1-3
                c0 = row.cells[0]
                r0 = c0.paragraphs[0].add_run(str(row_data[0]) if row_data[0] else "")
                r0.font.size = Pt(9)
                r0.font.color.rgb = grey_rgb
                merged = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
                rv = merged.paragraphs[0].add_run(str(row_data[1]) if row_data[1] else "—")
                rv.bold = True
                rv.font.size = Pt(9)
                rv.font.color.rgb = dark_rgb
        doc.add_paragraph()

    years_list = request.get('academic_years', [])
    years_str = ", ".join([f"{y.get('from_year')}-{y.get('to_year')}" for y in years_list]) if years_list else "—"
    status_disp = (request.get('enrollment_status') or "").replace('graduate','Graduate/Alumni').title()

    add_section_heading("Personal Information")
    add_field_table([
        ("First Name",     request.get('first_name',''),      "Last Name",       request.get('last_name','')),
        ("Middle Name",    request.get('middle_name') or "—", "Date of Birth",   request.get('date_of_birth','')),
        ("School ID",      request.get('school_id') or "—",   "Enrollment",      status_disp),
        ("Academic Years", years_str,                          "Last Form Class", request.get('last_form_class','')),
    ])

    add_section_heading("Contact Information")
    add_field_table([
        ("Wolmer's Email", request.get('wolmers_email') or request.get('student_email',''),
         "Personal Email", request.get('personal_email','')),
        ("Phone Number",   request.get('phone_number',''), "", ""),
    ])

    add_section_heading("External Examinations")
    p = doc.add_paragraph()
    p.add_run(format_external_exams(request.get('external_exams', []))).bold = True
    doc.add_paragraph()

    coll_map = {'pickup': 'Pickup at Bursary', 'emailed': 'Email to Institution', 'delivery': 'Courier Delivery (DHL)'}
    add_section_heading("Request Details")
    req_rows = [
        ("Reason", (request.get('reason','') or '').replace('_',' ').title(),
         "Number of Copies", str(request.get('number_of_copies') or 1)),
        ("Collection Method", coll_map.get(request.get('collection_method',''),''),
         "Received Before?", request.get('received_transcript_before') or "—"),
    ]
    if request.get('other_reason'):
        req_rows.append(("Other Reason", request.get('other_reason','')))
    if request.get('delivery_address'):
        req_rows.append(("Delivery Address", request.get('delivery_address','')))
    add_field_table(req_rows)

    add_section_heading("Destination Institution")
    add_field_table([
        ("Institution Name",  request.get('institution_name',''),    "Phone", request.get('institution_phone','')),
        ("Institution Email", request.get('institution_email',''),    "",      ""),
        ("Address", request.get('institution_address','')),
    ])

    # ---- ADMINISTRATIVE CLEARANCE ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Administrative Clearance")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = maroon_rgb

    ac = format_clearance_for_doc(request.get('administrative_clearance'))
    check_tbl = doc.add_table(rows=3, cols=2)
    check_tbl.style = 'Table Grid'

    def set_check_cell(cell, label, checked):
        p = cell.paragraphs[0]
        run_check = p.add_run('\u2611 ' if checked else '\u2610 ')
        try:   run_check.font.name = 'Segoe UI Symbol'
        except Exception: pass
        run_check.font.size = Pt(11)
        run_label = p.add_run(label)
        run_label.font.size = Pt(10)

    set_check_cell(check_tbl.cell(0, 0), "No fees outstanding", ac['no_fees'])
    set_check_cell(check_tbl.cell(0, 1), "No outstanding administrative obligations", ac['no_admin'])
    check_tbl.cell(1, 0).paragraphs[0].add_run(f"Amount Paid: {ac['amount']}").font.size = Pt(10)
    check_tbl.cell(1, 1).paragraphs[0].add_run(f"Receipt #: {ac['receipt']}     Date: {ac['date']}").font.size = Pt(10)
    proc_merged = check_tbl.cell(2, 0).merge(check_tbl.cell(2, 1))
    pby = f"Processed by: {ac['updated_by']}" if ac['updated_by'] != '—' else ""
    proc_merged.paragraphs[0].add_run(pby).font.size = Pt(9)

    # ---- STAFF NOTES ----
    if request.get('staff_notes'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(f"Staff Notes: {request['staff_notes']}")
        r.font.size = Pt(10)
        r.font.color.rgb = blue_rgb

    # ---- STATUS HISTORY ----
    timeline = request.get('timeline', [])
    if timeline:
        doc.add_paragraph()
        add_section_heading("Status History")
        tl_tbl = doc.add_table(rows=1+len(timeline), cols=4)
        tl_tbl.style = 'Table Grid'
        headers = ["Status", "Date", "Updated By", "Note"]
        for j, h in enumerate(headers):
            c = tl_tbl.cell(0, j)
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = maroon_rgb
        for i, entry in enumerate(timeline, start=1):
            row = tl_tbl.rows[i]
            vals = [
                entry.get('status',''),
                format_date_for_export(entry.get('timestamp','')),
                entry.get('updated_by','System'),
                entry.get('note','') or '—',
            ]
            for j, val in enumerate(vals):
                c = row.cells[j]
                r = c.paragraphs[0].add_run(str(val))
                r.font.size = Pt(9)
                r.font.color.rgb = dark_rgb

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_recommendation_pdf(request: dict) -> io.BytesIO:
    """Generate a formatted single-record recommendation letter PDF"""
    buffer = io.BytesIO()
    W = 495  # usable width
    doc_pdf = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    story = []
    styles = getSampleStyleSheet()

    maroon       = HexColor('#7b1e2c')
    gold_color   = HexColor('#b8860b')
    light_gold   = HexColor('#fffbeb')
    dark_text    = HexColor('#1c1917')
    mid_grey     = HexColor('#78716c')
    light_border = HexColor('#e7e5e4')
    red_bg       = HexColor('#fef2f2')
    blue_bg      = HexColor('#f0f9ff')

    title_style   = ParagraphStyle('Title',    parent=styles['Normal'], fontSize=16, textColor=maroon,     spaceAfter=4,  alignment=1, fontName='Helvetica-Bold')
    subtitle_style= ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=10, textColor=mid_grey,   spaceAfter=2,  alignment=1)
    section_style = ParagraphStyle('Section',  parent=styles['Normal'], fontSize=11, textColor=gold_color, fontName='Helvetica-Bold', spaceBefore=10, spaceAfter=4)
    label_style   = ParagraphStyle('Label',    parent=styles['Normal'], fontSize=9,  textColor=mid_grey,   leading=12, wordWrap='LTR')
    value_style   = ParagraphStyle('Value',    parent=styles['Normal'], fontSize=9,  textColor=dark_text,  fontName='Helvetica-Bold', leading=12, wordWrap='LTR')
    small_style   = ParagraphStyle('Small',    parent=styles['Normal'], fontSize=8,  textColor=mid_grey,   leading=11, wordWrap='LTR')
    check_style   = ParagraphStyle('Check',    parent=styles['Normal'], fontSize=9,  textColor=dark_text,  leading=12)

    LW  = 90
    VW  = 155
    COL4 = [LW, VW, LW, VW]

    # ---- HEADER ----
    story.append(Paragraph("New Recommendation Letter Form", title_style))
    story.append(Spacer(1, 6))

    try:
        import urllib.request as _ur
        logo_data = io.BytesIO(_ur.urlopen(WOLMERS_LOGO_URL, timeout=5).read())
        logo_img = RLImage(logo_data, width=50, height=50)
    except Exception:
        logo_img = None

    school_para_r = Paragraph(
        "<b>WOLMER'S BOYS' SCHOOL</b><br/>"
        "National Heroes Circle, Kingston 4<br/>"
        "Bursary: 876 922 4055 / 876 948 4807  |  wbs.bursary@wolmers.org<br/>"
        "Guidance: 876 922 8254  |  wbs.guidance@wolmers.org",
        ParagraphStyle('SchoolR', parent=styles['Normal'], fontSize=9, textColor=dark_text, leading=13, alignment=1, wordWrap='LTR')
    )
    if logo_img:
        header_table = Table([[logo_img], [school_para_r]], colWidths=[W])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('BOTTOMPADDING', (0,0), (0,0), 4),
            ('BOTTOMPADDING', (0,1), (0,1), 8),
            ('LINEBELOW', (0,1), (-1,1), 0.5, light_border),
        ]))
    else:
        header_table = Table([[school_para_r]], colWidths=[W])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('LINEBELOW', (0,0), (-1,0), 0.5, light_border),
        ]))
    story.append(header_table)
    story.append(Spacer(1, 10))

    status = request.get('status', 'Pending')
    story.append(Paragraph(
        f"Status: <b>{status}</b>  |  Submitted: <b>{format_date_for_export(request.get('created_at',''))}</b>"
        f"  |  Needed By: <b>{format_date_for_export(request.get('needed_by_date',''))}</b>",
        subtitle_style))
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=0.5, color=light_border))
    story.append(Spacer(1, 6))

    # ---- REJECTION BANNER ----
    if status == 'Rejected' and request.get('rejection_reason'):
        rej_box = Table([[
            Paragraph(f"<b>REJECTED</b> — {request['rejection_reason']}", ParagraphStyle(
                'RejNoteR', parent=styles['Normal'], fontSize=9, textColor=HexColor('#b91c1c'),
                leading=13, wordWrap='LTR'))
        ]], colWidths=[W])
        rej_box.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), red_bg),
            ('LINEBEFORE', (0,0), (0,-1), 3, HexColor('#ef4444')),
            ('TOPPADDING', (0,0), (-1,-1), 8), ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 12), ('RIGHTPADDING', (0,0), (-1,-1), 12),
        ]))
        story.append(rej_box)
        story.append(Spacer(1, 8))

    def section_header(text):
        story.append(Paragraph(text, section_style))
        story.append(HRFlowable(width="100%", thickness=0.3, color=light_border))
        story.append(Spacer(1, 4))

    def tc(l1, v1, l2="", v2=""):
        return {'l1': l1, 'v1': v1, 'l2': l2, 'v2': v2}

    def fc(label, value):
        return {'full': True, 'label': label, 'value': value}

    def add_info_block(rows):
        for r in rows:
            if r.get('full'):
                t = Table([[Paragraph(r['label'], label_style), Paragraph(r['value'] or '—', value_style)]], colWidths=[LW, W-LW])
            else:
                t = Table([[Paragraph(r['l1'], label_style), Paragraph(r.get('v1') or '—', value_style),
                            Paragraph(r.get('l2',''), label_style), Paragraph(r.get('v2') or '—', value_style)]], colWidths=COL4)
            t.setStyle(TableStyle([
                ('ROWBACKGROUNDS', (0,0), (-1,-1), [HexColor('#fffbeb'), HexColor('#fefce8')]),
                ('GRID', (0,0), (-1,-1), 0.3, light_border),
                ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ]))
            story.append(t)
        story.append(Spacer(1, 8))

    # ---- PERSONAL INFO ----
    section_header("Personal Information")
    years_list = request.get('years_attended', [])
    years_str = ", ".join([f"{y.get('from_year')}-{y.get('to_year')}" for y in years_list]) if years_list else request.get('years_attended_str', '—')
    add_info_block([
        tc("First Name",   request.get('first_name'),          "Last Name",   request.get('last_name')),
        tc("Middle Name",  request.get('middle_name') or "—",  "Date of Birth", request.get('date_of_birth')),
        tc("Wolmer's Email", request.get('student_email'),     "Personal Email", request.get('email')),
        tc("Phone",        request.get('phone_number'),         "Address",     request.get('address')),
        tc("Years Attended", years_str,                         "Last Form Class", request.get('last_form_class')),
        tc("Enrollment",   request.get('enrollment_status'),   "", ""),
    ])

    if request.get('co_curricular_activities'):
        section_header("Co-curricular Activities / Positions of Responsibility")
        story.append(Paragraph(request['co_curricular_activities'], value_style))
        story.append(Spacer(1, 8))

    section_header("External Examinations")
    story.append(Paragraph(format_external_exams(request.get('external_exams', [])), value_style))
    story.append(Spacer(1, 8))

    coll_map = {'pickup': 'Pickup at School', 'emailed': 'Email to Institution', 'delivery': 'Courier Delivery (DHL)'}
    section_header("Request Details")
    req_rows = [tc("Reason", (request.get('reason','') or '').replace('_',' ').title(),
                   "Collection", coll_map.get(request.get('collection_method',''), ''))]
    if request.get('other_reason'):
        req_rows.append(fc("Other Reason", request['other_reason']))
    if request.get('delivery_address'):
        req_rows.append(fc("Delivery Address", request['delivery_address']))
    add_info_block(req_rows)

    section_header("Destination Institution")
    add_info_block([
        tc("Institution", request.get('institution_name'), "Program", request.get('program_name')),
        tc("Directed To", request.get('directed_to') or "—", "", ""),
        fc("Address", request.get('institution_address')),
    ])

    # ---- ADMINISTRATIVE CLEARANCE ----
    story.append(HRFlowable(width="100%", thickness=1, color=gold_color))
    story.append(Spacer(1, 4))
    story.append(Paragraph("Administrative Clearance", ParagraphStyle('ClearHeader', parent=styles['Normal'],
        fontSize=12, textColor=gold_color, fontName='Helvetica-Bold', spaceBefore=4, spaceAfter=6)))

    ac = format_clearance_for_doc(request.get('administrative_clearance'))

    def _draw_checkbox_gold(checked: bool):
        from reportlab.graphics.shapes import Drawing, Rect, Line
        sz = 8
        d = Drawing(sz + 4, sz + 4)
        d.add(Rect(2, 2, sz, sz, strokeColor=HexColor('#b8860b'), strokeWidth=0.8, fillColor=HexColor('#ffffff')))
        if checked:
            d.add(Line(3, 6, 5.5, 3.5, strokeColor=HexColor('#b8860b'), strokeWidth=1.2))
            d.add(Line(5.5, 3.5, 10, 9, strokeColor=HexColor('#b8860b'), strokeWidth=1.2))
        return d

    ct_gold = Table([
        [_draw_checkbox_gold(ac['no_fees']),  Paragraph("<b>No fees outstanding</b>", check_style)],
        [_draw_checkbox_gold(ac['no_admin']), Paragraph("<b>No outstanding administrative obligations</b>", check_style)],
    ], colWidths=[16, W-16])
    ct_gold.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), light_gold),
        ('GRID', (0,0), (-1,-1), 0, light_gold),
        ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8), ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(ct_gold)

    pay_label = ParagraphStyle('PayLabelG', parent=styles['Normal'], fontSize=9, textColor=HexColor('#78716c'), wordWrap='LTR')
    pay_val   = ParagraphStyle('PayValG',   parent=styles['Normal'], fontSize=9, textColor=HexColor('#1c1917'), fontName='Helvetica-Bold', wordWrap='LTR')
    pay_rows = [[
        Paragraph("Amount Paid:", pay_label), Paragraph(ac['amount'], pay_val),
        Paragraph("Receipt #:", pay_label),   Paragraph(ac['receipt'], pay_val),
        Paragraph("Date:", pay_label),        Paragraph(ac['date'], pay_val),
    ]]
    if ac['updated_by'] != '—':
        pay_rows.append([Paragraph("Processed by:", pay_label), Paragraph(ac['updated_by'], pay_val), '', '', '', ''])
    pt_gold = Table(pay_rows, colWidths=[65, 90, 55, 110, 40, W-360])
    pt_gold.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), light_gold),
        ('GRID', (0,0), (-1,-1), 0, light_gold),
        ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8), ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('SPAN', (1,1), (-1,1)) if ac['updated_by'] != '—' else ('SPAN', (0,0), (0,0)),
    ]))
    story.append(pt_gold)
    story.append(Spacer(1, 10))

    # ---- STAFF NOTES ----
    if request.get('staff_notes'):
        notes_box = Table([[
            Paragraph(f"<b>Staff Notes:</b> {request['staff_notes']}", ParagraphStyle(
                'StaffNoteR', parent=styles['Normal'], fontSize=9, textColor=HexColor('#0369a1'),
                leading=13, wordWrap='LTR'))
        ]], colWidths=[W])
        notes_box.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), blue_bg),
            ('LINEBEFORE', (0,0), (0,-1), 3, HexColor('#0284c7')),
            ('TOPPADDING', (0,0), (-1,-1), 7), ('BOTTOMPADDING', (0,0), (-1,-1), 7),
            ('LEFTPADDING', (0,0), (-1,-1), 12), ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(notes_box)
        story.append(Spacer(1, 8))

    # ---- STATUS HISTORY ----
    timeline = request.get('timeline', [])
    if timeline:
        story.append(HRFlowable(width="100%", thickness=0.5, color=light_border))
        story.append(Spacer(1, 4))
        story.append(Paragraph("Status History", ParagraphStyle('TLHeaderR', parent=styles['Normal'],
            fontSize=11, textColor=gold_color, fontName='Helvetica-Bold', spaceBefore=4, spaceAfter=4)))
        tl_data = [[
            Paragraph("<b>Status</b>", label_style), Paragraph("<b>Date</b>", label_style),
            Paragraph("<b>Updated By</b>", label_style), Paragraph("<b>Note</b>", label_style),
        ]]
        for entry in timeline:
            tl_data.append([
                Paragraph(entry.get('status',''), value_style),
                Paragraph(format_date_for_export(entry.get('timestamp','')), small_style),
                Paragraph(entry.get('updated_by','System'), small_style),
                Paragraph(entry.get('note','') or '—', small_style),
            ])
        tl_t = Table(tl_data, colWidths=[90, 80, 120, W-290])
        tl_t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), gold_color),
            ('TEXTCOLOR', (0,0), (-1,0), HexColor('#ffffff')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [HexColor('#fffbeb'), HexColor('#fefce8')]),
            ('GRID', (0,0), (-1,-1), 0.3, light_border),
            ('TOPPADDING', (0,0), (-1,-1), 4), ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        story.append(tl_t)

    doc_pdf.build(story)
    buffer.seek(0)
    return buffer

    maroon = HexColor('#7b1e2c')
    gold_color = HexColor('#b8860b')
    light_gold = HexColor('#fffbeb')
    dark_text = HexColor('#1c1917')
    mid_grey = HexColor('#78716c')
    light_border = HexColor('#e7e5e4')

    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, textColor=maroon, spaceAfter=4, alignment=1, fontName='Helvetica-Bold')
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=10, textColor=mid_grey, spaceAfter=2, alignment=1)
    section_style = ParagraphStyle('Section', parent=styles['Normal'], fontSize=11, textColor=gold_color, fontName='Helvetica-Bold', spaceBefore=10, spaceAfter=4)
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontSize=9, textColor=mid_grey)
    value_style = ParagraphStyle('Value', parent=styles['Normal'], fontSize=10, textColor=dark_text, fontName='Helvetica-Bold')
    small_style = ParagraphStyle('Small', parent=styles['Normal'], fontSize=8, textColor=mid_grey)
    check_style = ParagraphStyle('Check', parent=styles['Normal'], fontSize=10, textColor=dark_text)

    # ---- HEADER ----
    story.append(Paragraph("New Recommendation Letter Form", title_style))
    story.append(Spacer(1, 6))

    try:
        import urllib.request
        logo_data = io.BytesIO(urllib.request.urlopen(WOLMERS_LOGO_URL, timeout=5).read())
        logo_img = RLImage(logo_data, width=55, height=55)
    except Exception:
        logo_img = Paragraph("", styles['Normal'])

    school_para = Paragraph(
        "<b>WOLMER'S BOYS' SCHOOL</b><br/>"
        "National Heroes Circle, Kingston 4",
        ParagraphStyle('School', parent=styles['Normal'], fontSize=10, textColor=dark_text, leading=16, alignment=1)
    )

    header_data = [[logo_img, school_para]]
    header_table = Table(header_data, colWidths=[70, 400])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
        ('LINEBELOW', (0, 0), (-1, 0), 0.5, light_border),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 14))

    story.append(Paragraph(f"Status: <b>{request.get('status','')}</b>  |  Submitted: <b>{format_date_for_export(request.get('created_at',''))}</b>  |  Needed By: <b>{format_date_for_export(request.get('needed_by_date',''))}</b>", subtitle_style))
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=0.5, color=light_border))
    story.append(Spacer(1, 8))

    def section_header(text):
        story.append(Paragraph(text, section_style))
        story.append(HRFlowable(width="100%", thickness=0.3, color=light_border))
        story.append(Spacer(1, 4))

    def two_col_row(l1, v1, l2="", v2=""):
        return [[Paragraph(l1, label_style), Paragraph(v1 or "—", value_style), Paragraph(l2, label_style), Paragraph(v2 or "—", value_style)]]

    def make_table(rows):
        t = Table(rows, colWidths=[95, 160, 95, 120])
        t.setStyle(TableStyle([
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [HexColor('#fffbeb'), HexColor('#fefce8')]),
            ('GRID', (0, 0), (-1, -1), 0.3, light_border),
            ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        return t

    # ---- PERSONAL INFO ----
    section_header("Personal Information")
    years_list = request.get('years_attended', [])
    years_str = ", ".join([f"{y.get('from_year')}-{y.get('to_year')}" for y in years_list]) if years_list else request.get('years_attended_str', '—')
    rows = []
    rows += two_col_row("First Name", request.get('first_name'), "Last Name", request.get('last_name'))
    rows += two_col_row("Middle Name", request.get('middle_name') or "—", "Date of Birth", request.get('date_of_birth'))
    rows += two_col_row("Wolmer's Email", request.get('student_email'), "Personal Email", request.get('email'))
    rows += two_col_row("Phone", request.get('phone_number'), "Address", request.get('address'))
    rows += two_col_row("Years Attended", years_str, "Last Form Class", request.get('last_form_class'))
    rows += two_col_row("Enrollment Status", request.get('enrollment_status'), "", "")
    story.append(make_table(rows))
    story.append(Spacer(1, 8))

    if request.get('co_curricular_activities'):
        section_header("Co-curricular Activities / Positions of Responsibility")
        story.append(Paragraph(request['co_curricular_activities'], value_style))
        story.append(Spacer(1, 8))

    # ---- EXTERNAL EXAMS ----
    section_header("External Examinations")
    story.append(Paragraph(format_external_exams(request.get('external_exams', [])), value_style))
    story.append(Spacer(1, 8))

    # ---- REQUEST DETAILS ----
    section_header("Request Details")
    coll_map = {'pickup': 'Pickup at School', 'emailed': 'Email to Institution', 'delivery': 'Courier Delivery (DHL)'}
    rows3 = []
    rows3 += two_col_row("Reason", request.get('reason'), "Collection Method", coll_map.get(request.get('collection_method',''), ''))
    if request.get('other_reason'):
        rows3 += [[Paragraph("Other Reason", label_style), Paragraph(request.get('other_reason',''), value_style), '', '']]
    if request.get('delivery_address'):
        rows3 += [[Paragraph("Delivery Address", label_style), Paragraph(request.get('delivery_address',''), value_style), '', '']]
    story.append(make_table(rows3))
    story.append(Spacer(1, 8))

    # ---- INSTITUTION ----
    section_header("Destination Institution")
    rows4 = []
    rows4 += two_col_row("Institution Name", request.get('institution_name'), "Program", request.get('program_name'))
    rows4 += two_col_row("Directed To", request.get('directed_to') or "—", "", "")
    rows4 += [[Paragraph("Address", label_style), Paragraph(request.get('institution_address',''), value_style), '', '']]
    story.append(make_table(rows4))
    story.append(Spacer(1, 12))

    # ---- ADMINISTRATIVE CLEARANCE ----
    story.append(HRFlowable(width="100%", thickness=1, color=gold_color))
    story.append(Spacer(1, 4))
    story.append(Paragraph("Administrative Clearance", ParagraphStyle('ClearHeader', parent=styles['Normal'], fontSize=12, textColor=gold_color, fontName='Helvetica-Bold', spaceBefore=4, spaceAfter=6)))

    ac = format_clearance_for_doc(request.get('administrative_clearance'))

    def _draw_checkbox_gold(checked: bool):
        from reportlab.graphics.shapes import Drawing, Rect, Line
        sz = 8
        d = Drawing(sz + 4, sz + 4)
        d.add(Rect(2, 2, sz, sz, strokeColor=HexColor('#b8860b'), strokeWidth=0.8, fillColor=HexColor('#ffffff')))
        if checked:
            d.add(Line(3, 6, 5.5, 3.5, strokeColor=HexColor('#b8860b'), strokeWidth=1.2))
            d.add(Line(5.5, 3.5, 10, 9, strokeColor=HexColor('#b8860b'), strokeWidth=1.2))
        return d

    def checkbox_row_gold(label: str, checked: bool) -> list:
        return [_draw_checkbox_gold(checked), Paragraph(f"<b>{label}</b>", check_style)]

    clearance_rows_gold = [
        checkbox_row_gold("No fees outstanding", ac['no_fees']),
        checkbox_row_gold("No outstanding administrative obligations", ac['no_admin']),
    ]
    ct_gold = Table(clearance_rows_gold, colWidths=[16, 454])
    ct_gold.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), light_gold),
        ('GRID', (0, 0), (-1, -1), 0, light_gold),
        ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(ct_gold)

    pay_label = ParagraphStyle('PayLabelG', parent=styles['Normal'], fontSize=9, textColor=HexColor('#78716c'))
    pay_val = ParagraphStyle('PayValG', parent=styles['Normal'], fontSize=9, textColor=HexColor('#1c1917'), fontName='Helvetica-Bold')
    payment_rows_gold = [[
        Paragraph("Amount Paid:", pay_label), Paragraph(ac['amount'], pay_val),
        Paragraph("Receipt #:", pay_label), Paragraph(ac['receipt'], pay_val),
        Paragraph("Date:", pay_label), Paragraph(ac['date'], pay_val),
    ]]
    if ac['updated_by'] != '—':
        payment_rows_gold.append([
            Paragraph("Processed by:", pay_label), Paragraph(ac['updated_by'], pay_val),
            '', '', '', ''
        ])
    pt_gold = Table(payment_rows_gold, colWidths=[60, 85, 55, 90, 35, 145])
    pt_gold.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), light_gold),
        ('GRID', (0, 0), (-1, -1), 0, light_gold),
        ('TOPPADDING', (0, 0), (-1, -1), 5), ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8), ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(pt_gold)

    doc_pdf.build(story)
    buffer.seek(0)
    return buffer


def build_recommendation_docx(request: dict) -> io.BytesIO:
    """Generate a formatted single-record recommendation letter DOCX"""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import urllib.request

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    maroon_rgb = RGBColor(0x7b, 0x1e, 0x2c)
    gold_rgb   = RGBColor(0xb8, 0x86, 0x0b)
    grey_rgb   = RGBColor(0x78, 0x71, 0x6c)
    dark_rgb   = RGBColor(0x1c, 0x19, 0x17)
    red_rgb    = RGBColor(0xb9, 0x1c, 0x1c)
    blue_rgb   = RGBColor(0x03, 0x69, 0xa1)

    # ---- HEADER ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("New Recommendation Letter Form")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = maroon_rgb

    # Logo centred on its own line
    try:
        logo_bytes = io.BytesIO(urllib.request.urlopen(WOLMERS_LOGO_URL, timeout=5).read())
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_bytes, width=Inches(0.75))
    except Exception:
        pass

    # School info centred below logo
    school_lines = [
        ("WOLMER'S BOYS' SCHOOL", True,  Pt(12), maroon_rgb),
        ("National Heroes Circle, Kingston 4", False, Pt(9), dark_rgb),
        ("Bursary: 876 922 4055  |  wbs.bursary@wolmers.org", False, Pt(9), grey_rgb),
        ("Guidance: 876 922 8254  |  wbs.guidance@wolmers.org", False, Pt(9), grey_rgb),
    ]
    for text, bold, size, color in school_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = size
        r.font.color.rgb = color

    # Status line
    status_para = doc.add_paragraph()
    status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = request.get('status', 'Pending')
    sr = status_para.add_run(
        f"Status: {st}  |  Submitted: {format_date_for_export(request.get('created_at',''))}"
        f"  |  Needed By: {format_date_for_export(request.get('needed_by_date',''))}"
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = grey_rgb

    # Rejection banner
    if st == 'Rejected' and request.get('rejection_reason'):
        rej_p = doc.add_paragraph()
        rej_r = rej_p.add_run(f"REJECTED — {request['rejection_reason']}")
        rej_r.bold = True
        rej_r.font.size = Pt(10)
        rej_r.font.color.rgb = red_rgb

    doc.add_paragraph()

    def add_section_heading(text, color=gold_rgb):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = color

    def add_field_table(rows_data):
        tbl = doc.add_table(rows=len(rows_data), cols=4)
        tbl.style = 'Table Grid'
        for i, row_data in enumerate(rows_data):
            row = tbl.rows[i]
            if len(row_data) >= 4:
                for j, (cell_text, is_bold) in enumerate(zip(row_data, [False, True, False, True])):
                    cell = row.cells[j]
                    p = cell.paragraphs[0]
                    r = p.add_run(str(cell_text) if cell_text else "—")
                    r.bold = is_bold
                    r.font.size = Pt(9)
                    r.font.color.rgb = dark_rgb if is_bold else grey_rgb
            else:
                c0 = row.cells[0]
                r0 = c0.paragraphs[0].add_run(str(row_data[0]) if row_data[0] else "")
                r0.font.size = Pt(9)
                r0.font.color.rgb = grey_rgb
                merged = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
                rv = merged.paragraphs[0].add_run(str(row_data[1]) if row_data[1] else "—")
                rv.bold = True
                rv.font.size = Pt(9)
                rv.font.color.rgb = dark_rgb
        doc.add_paragraph()

    years_list = request.get('years_attended', [])
    years_str = ", ".join([f"{y.get('from_year')}-{y.get('to_year')}" for y in years_list]) if years_list else request.get('years_attended_str', '—')

    add_section_heading("Personal Information")
    add_field_table([
        ("First Name",     request.get('first_name',''),      "Last Name",       request.get('last_name','')),
        ("Middle Name",    request.get('middle_name') or "—", "Date of Birth",   request.get('date_of_birth','')),
        ("Wolmer's Email", request.get('student_email',''),   "Personal Email",  request.get('email','')),
        ("Phone",          request.get('phone_number',''),    "Address",         request.get('address','')),
        ("Years Attended", years_str,                          "Last Form Class", request.get('last_form_class','')),
        ("Enrollment",     request.get('enrollment_status',''), "", ""),
    ])

    if request.get('co_curricular_activities'):
        add_section_heading("Co-curricular Activities")
        p = doc.add_paragraph()
        p.add_run(request['co_curricular_activities']).bold = True
        doc.add_paragraph()

    add_section_heading("External Examinations")
    p = doc.add_paragraph()
    p.add_run(format_external_exams(request.get('external_exams', []))).bold = True
    doc.add_paragraph()

    coll_map = {'pickup': 'Pickup at School', 'emailed': 'Email to Institution', 'delivery': 'Courier Delivery (DHL)'}
    add_section_heading("Request Details")
    req_rows = [
        ("Reason", (request.get('reason','') or '').replace('_',' ').title(),
         "Collection Method", coll_map.get(request.get('collection_method',''),'')),
    ]
    if request.get('delivery_address'):
        req_rows.append(("Delivery Address", request.get('delivery_address','')))
    add_field_table(req_rows)

    add_section_heading("Destination Institution")
    add_field_table([
        ("Institution Name", request.get('institution_name',''), "Program", request.get('program_name','')),
        ("Directed To", request.get('directed_to') or "—", "", ""),
        ("Address", request.get('institution_address','')),
    ])

    # ---- ADMINISTRATIVE CLEARANCE ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Administrative Clearance")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = gold_rgb

    ac = format_clearance_for_doc(request.get('administrative_clearance'))
    check_tbl = doc.add_table(rows=3, cols=2)
    check_tbl.style = 'Table Grid'

    def set_check_cell_gold(cell, label, checked):
        p = cell.paragraphs[0]
        run_check = p.add_run('\u2611 ' if checked else '\u2610 ')
        try:   run_check.font.name = 'Segoe UI Symbol'
        except Exception: pass
        run_check.font.size = Pt(11)
        run_label = p.add_run(label)
        run_label.font.size = Pt(10)

    set_check_cell_gold(check_tbl.cell(0, 0), "No fees outstanding", ac['no_fees'])
    set_check_cell_gold(check_tbl.cell(0, 1), "No outstanding administrative obligations", ac['no_admin'])
    check_tbl.cell(1, 0).paragraphs[0].add_run(f"Amount Paid: {ac['amount']}").font.size = Pt(10)
    check_tbl.cell(1, 1).paragraphs[0].add_run(f"Receipt #: {ac['receipt']}     Date: {ac['date']}").font.size = Pt(10)
    proc_merged = check_tbl.cell(2, 0).merge(check_tbl.cell(2, 1))
    pby = f"Processed by: {ac['updated_by']}" if ac['updated_by'] != '—' else ""
    proc_merged.paragraphs[0].add_run(pby).font.size = Pt(9)

    # ---- STAFF NOTES ----
    if request.get('staff_notes'):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(f"Staff Notes: {request['staff_notes']}")
        r.font.size = Pt(10)
        r.font.color.rgb = blue_rgb

    # ---- STATUS HISTORY ----
    timeline = request.get('timeline', [])
    if timeline:
        doc.add_paragraph()
        add_section_heading("Status History", gold_rgb)
        tl_tbl = doc.add_table(rows=1+len(timeline), cols=4)
        tl_tbl.style = 'Table Grid'
        for j, h in enumerate(["Status", "Date", "Updated By", "Note"]):
            c = tl_tbl.cell(0, j)
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = gold_rgb
        for i, entry in enumerate(timeline, start=1):
            row = tl_tbl.rows[i]
            for j, val in enumerate([
                entry.get('status',''),
                format_date_for_export(entry.get('timestamp','')),
                entry.get('updated_by','System'),
                entry.get('note','') or '—',
            ]):
                row.cells[j].paragraphs[0].add_run(str(val)).font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

    maroon_rgb = RGBColor(0x7b, 0x1e, 0x2c)
    gold_rgb = RGBColor(0xb8, 0x86, 0x0b)
    grey_rgb = RGBColor(0x78, 0x71, 0x6c)
    dark_rgb = RGBColor(0x1c, 0x19, 0x17)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("New Recommendation Letter Form")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = maroon_rgb

    # Header table - logo + school info
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_cell = header_tbl.cell(0, 0)
    logo_cell.width = Inches(1.0)
    try:
        logo_bytes = io.BytesIO(urllib.request.urlopen(WOLMERS_LOGO_URL, timeout=5).read())
        para = logo_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(logo_bytes, width=Inches(0.85))
    except Exception:
        logo_cell.paragraphs[0].add_run("WBS")

    info_cell = header_tbl.cell(0, 1)
    school_lines = [
        ("WOLMER'S BOYS' SCHOOL", True, Pt(13), maroon_rgb),
        ("National Heroes Circle, Kingston 4", False, Pt(10), dark_rgb),
    ]
    first = True
    for text, bold, size, color in school_lines:
        p = info_cell.paragraphs[0] if first else info_cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = size
        r.font.color.rgb = color

    doc.add_paragraph()

    def add_section_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = gold_rgb

    def add_field_table(rows_data):
        tbl = doc.add_table(rows=len(rows_data), cols=4)
        tbl.style = 'Table Grid'
        for i, row_data in enumerate(rows_data):
            row = tbl.rows[i]
            for j, (cell_text, is_bold) in enumerate(zip(row_data, [False, True, False, True])):
                cell = row.cells[j]
                p = cell.paragraphs[0]
                r = p.add_run(str(cell_text) if cell_text else "—")
                r.bold = is_bold
                r.font.size = Pt(9)
                r.font.color.rgb = dark_rgb if is_bold else grey_rgb
        doc.add_paragraph()

    years_list = request.get('years_attended', [])
    years_str = ", ".join([f"{y.get('from_year')}-{y.get('to_year')}" for y in years_list]) if years_list else request.get('years_attended_str', '—')

    add_section_heading("Personal Information")
    add_field_table([
        ("First Name", request.get('first_name',''), "Last Name", request.get('last_name','')),
        ("Middle Name", request.get('middle_name') or "—", "Date of Birth", request.get('date_of_birth','')),
        ("Wolmer's Email", request.get('student_email',''), "Personal Email", request.get('email','')),
        ("Phone", request.get('phone_number',''), "Address", request.get('address','')),
        ("Years Attended", years_str, "Last Form Class", request.get('last_form_class','')),
        ("Enrollment Status", request.get('enrollment_status',''), "", ""),
    ])

    if request.get('co_curricular_activities'):
        add_section_heading("Co-curricular Activities")
        p = doc.add_paragraph()
        p.add_run(request['co_curricular_activities']).bold = True
        doc.add_paragraph()

    add_section_heading("External Examinations")
    p = doc.add_paragraph()
    p.add_run(format_external_exams(request.get('external_exams', []))).bold = True
    doc.add_paragraph()

    add_section_heading("Request Details")
    coll_map = {'pickup': 'Pickup at School', 'emailed': 'Email to Institution', 'delivery': 'Courier Delivery (DHL)'}
    req_rows = [
        ("Reason", request.get('reason',''), "Collection Method", coll_map.get(request.get('collection_method',''),'')),
    ]
    if request.get('delivery_address'):
        req_rows.append(("Delivery Address", request.get('delivery_address',''), "", ""))
    add_field_table(req_rows)

    add_section_heading("Destination Institution")
    add_field_table([
        ("Institution Name", request.get('institution_name',''), "Program", request.get('program_name','')),
        ("Directed To", request.get('directed_to') or "—", "", ""),
        ("Address", request.get('institution_address',''), "", ""),
    ])

    # Administrative Clearance
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Administrative Clearance")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = gold_rgb

    ac = format_clearance_for_doc(request.get('administrative_clearance'))
    check_tbl = doc.add_table(rows=2, cols=2)
    check_tbl.style = 'Table Grid'

    def set_check_cell_gold(cell, label, checked):
        p = cell.paragraphs[0]
        run_check = p.add_run('\u2611 ' if checked else '\u2610 ')
        try:
            run_check.font.name = 'Segoe UI Symbol'
        except Exception:
            pass
        run_check.font.size = Pt(11)
        run_label = p.add_run(label)
        run_label.font.size = Pt(10)

    set_check_cell_gold(check_tbl.cell(0, 0), "No fees outstanding", ac['no_fees'])
    set_check_cell_gold(check_tbl.cell(0, 1), "No outstanding administrative obligations", ac['no_admin'])
    check_tbl.cell(1, 0).paragraphs[0].add_run(f"Amount Paid: {ac['amount']}").font.size = Pt(10)
    check_tbl.cell(1, 1).paragraphs[0].add_run(f"Receipt #: {ac['receipt']}     Date: {ac['date']}").font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@api_router.get("/requests/{request_id}/export/{format}")
async def export_single_transcript(
    request_id: str,
    format: str,
    current_user: dict = Depends(get_current_user)
):
    """Export a single transcript request as PDF or DOCX"""
    if current_user.get('role') not in ['admin', 'staff']:
        raise HTTPException(status_code=403, detail="Only admin/staff can export requests")
    if format not in ['pdf', 'docx']:
        raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")

    request_data = await db.transcript_requests.find_one({"id": request_id})
    if not request_data:
        raise HTTPException(status_code=404, detail="Request not found")

    request_data = normalize_transcript_data(request_data)
    short_id = request_id[:8]

    if format == 'pdf':
        buffer = build_transcript_pdf(request_data)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=transcript_request_{short_id}.pdf"}
        )
    else:
        buffer = build_transcript_docx(request_data)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=transcript_request_{short_id}.docx"}
        )


@api_router.get("/recommendations/{request_id}/export/{format}")
async def export_single_recommendation(
    request_id: str,
    format: str,
    current_user: dict = Depends(get_current_user)
):
    """Export a single recommendation request as PDF or DOCX"""
    if current_user.get('role') not in ['admin', 'staff']:
        raise HTTPException(status_code=403, detail="Only admin/staff can export requests")
    if format not in ['pdf', 'docx']:
        raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")

    request_data = await db.recommendation_requests.find_one({"id": request_id})
    if not request_data:
        raise HTTPException(status_code=404, detail="Recommendation request not found")

    request_data = normalize_recommendation_data(request_data)
    short_id = request_id[:8]

    if format == 'pdf':
        buffer = build_recommendation_pdf(request_data)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=recommendation_request_{short_id}.pdf"}
        )
    else:
        buffer = build_recommendation_docx(request_data)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=recommendation_request_{short_id}.docx"}
        )


# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
